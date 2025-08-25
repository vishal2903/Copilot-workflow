#!/usr/bin/env python3
"""
Google Drive LLM Analyzer - Single File Version
A Python tool that uses OpenAI GPT-4o-mini to analyze Google Drive documents and generate reports.
"""

import os
import sys
import click
import io
import sys
import json
import click
import PyPDF2
from datetime import datetime
import pytz
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
import logging
from concurrent.futures import ThreadPoolExecutor, as_completed
import time
import asyncio
from dotenv import load_dotenv
# LiveKit imports for voice functionality

from livekit import agents
from livekit.agents import AgentSession, Agent, RoomInputOptions, function_tool, RunContext
from livekit.plugins import openai as lk_openai, noise_cancellation, silero
from livekit.agents import cli as lk_cli
from livekit.agents.worker import WorkerOptions 

# Rich imports for beautiful CLI
from rich.console import Console
from rich.table import Table
from rich.panel import Panel
from rich.progress import Progress, SpinnerColumn, TextColumn
from rich.prompt import Confirm

# Google Drive API imports
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
import httplib2
from google_auth_httplib2 import AuthorizedHttp
import traceback
import logging
# OpenAI imports
from openai import OpenAI
from dotenv import load_dotenv
import tiktoken
import re
import urllib.parse
from docx import Document as DocxDocument
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
# ---------- LiveKit imports + availability flags (paste near top of file) ----------
# Make LiveKit imports robust across versions and set availability flags so the
# rest of the module can check safely for the library.
LIVEKIT_AVAILABLE = False
lk_cli = None
WorkerOptions = None
lk_openai = None
noise_cancellation = None
silero = None

try:
    # core public API
    from livekit.agents import cli as lk_cli
    # WorkerOptions sometimes lives at livekit.agents or livekit.agents.worker
    try:
        from livekit.agents import WorkerOptions
    except Exception:
        try:
            from livekit.agents.worker import WorkerOptions
        except Exception:
            WorkerOptions = None

    # plugins (may be separate packages depending on your env)
    try:
        from livekit.plugins import openai as lk_openai
    except Exception:
        lk_openai = None

    try:
        from livekit.plugins import noise_cancellation
    except Exception:
        noise_cancellation = None

    try:
        from livekit.plugins import silero
    except Exception:
        silero = None

    # mark available if at least the cli runner exists and WorkerOptions is present
    if lk_cli is not None and WorkerOptions is not None:
        LIVEKIT_AVAILABLE = True
except Exception:
    # keep defaults (None/False) — voice command will raise a helpful error later
    LIVEKIT_AVAILABLE = False
# -----------------------------------------------------------------------------

# Load .env early so LIVEKIT_* is available to the LiveKit runner
load_dotenv()


# =============================================================================
# CONFIGURATION AND LOGGING
# =============================================================================
_SURROGATE_HIGH = re.compile(r'[\uD800-\uDBFF](?![\uDC00-\uDFFF])')
_SURROGATE_LOW  = re.compile(r'(?<![\uD800-\uDBFF])[\uDC00-\uDFFF]')

def sanitize_string(s: str) -> str:
    """
    Convert invalid unicode/surrogates to U+FFFD replacement char.
    Also normalizes any other invalid sequences via utf-8 roundtrip if needed.
    """
    if not isinstance(s, str):
        return s
    # Replace lone surrogates
    s = _SURROGATE_HIGH.sub('\uFFFD', s)
    s = _SURROGATE_LOW.sub('\uFFFD', s)
    # Safe utf-8 roundtrip to remove any other invalid sequences
    try:
        return s.encode('utf-8', errors='replace').decode('utf-8', errors='replace')
    except Exception:
        return s

def sanitize(obj):
    if isinstance(obj, str):
        return sanitize_string(obj)
    if isinstance(obj, list):
        return [sanitize(x) for x in obj]
    if isinstance(obj, dict):
        return {k: sanitize(v) for k, v in obj.items()}
    return obj

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('gdrive_analyzer.log'),
        logging.StreamHandler()
    ]
)
logging.getLogger("livekit.agents").setLevel(logging.DEBUG)
logger = logging.getLogger(__name__)

# Set more verbose logging for specific modules when debugging
logging.getLogger('googleapiclient.discovery').setLevel(logging.WARNING)
logging.getLogger('google.auth').setLevel(logging.WARNING)
logging.getLogger('urllib3.connectionpool').setLevel(logging.WARNING)
# At the top of your file, add:
logging.getLogger("livekit").setLevel(logging.DEBUG)
logging.getLogger("livekit.agents").setLevel(logging.DEBUG)



class GoogleDocsFormatter:
    """Format content for Google Docs without markdown syntax"""
    
    @staticmethod
    def format_content_for_google_docs(content: str) -> str:
        """Remove markdown and format content for clean Google Docs display"""
        # Remove markdown headers and replace with clean titles
        content = re.sub(r'^# (.+)$', r'\1', content, flags=re.MULTILINE)
        content = re.sub(r'^## (.+)$', r'\1', content, flags=re.MULTILINE)  
        content = re.sub(r'^### (.+)$', r'\1', content, flags=re.MULTILINE)
        
        # Remove markdown bold/italic syntax
        content = re.sub(r'\*\*(.+?)\*\*', r'\1', content)
        content = re.sub(r'\*(.+?)\*', r'\1', content)
        
        # Clean up bullet points - remove markdown syntax
        content = re.sub(r'^[\s]*[-*+]\s+', '• ', content, flags=re.MULTILINE)
        
        # Remove markdown links syntax but keep the URLs
        content = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'\1 (\2)', content)
        
        return content.strip()
    
    @staticmethod
    def create_formatted_requests(content: str):
        """Create Google Docs API requests for properly formatted content"""
        requests = []
        lines = content.split('\n')
        current_index = 1
        
        for line in lines:
            if line.strip():
                # Detect headings and format accordingly
                if line.strip() and not line.startswith('•') and not line.startswith('-'):
                    # Potential heading - make it bold and larger
                    requests.append({
                        'insertText': {
                            'location': {'index': current_index},
                            'text': line + '\n'
                        }
                    })
                    
                    # Make it bold and larger for heading
                    requests.append({
                        'updateTextStyle': {
                            'range': {
                                'startIndex': current_index,
                                'endIndex': current_index + len(line)
                            },
                            'textStyle': {
                                'bold': True,
                                'fontSize': {'magnitude': 14, 'unit': 'PT'}
                            },
                            'fields': 'bold,fontSize'
                        }
                    })
                    current_index += len(line) + 1
                else:
                    # Regular content
                    requests.append({
                        'insertText': {
                            'location': {'index': current_index},
                            'text': line + '\n'
                        }
                    })
                    current_index += len(line) + 1
            else:
                # Empty line - add spacing
                requests.append({
                    'insertText': {
                        'location': {'index': current_index},
                        'text': '\n'
                    }
                })
                current_index += 1
        
        return requests

class DocxReportGenerator:
    """Generate .docx reports for local storage"""
    
    def __init__(self, downloads_dir: str):
        self.downloads_dir = downloads_dir
    
    def create_docx_report(self, content: str, title: str) -> str:
        """Create a professionally formatted .docx report"""
        doc = DocxDocument()
        
        # Add title
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Process content line by line
        lines = content.split('\n')
        
        for line in lines:
            if line.strip():
                # Detect different content types
                if line.startswith('# '):
                    # Main heading
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    # Sub heading
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    # Sub-sub heading  
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('• ') or line.startswith('- '):
                    # Bullet point
                    para = doc.add_paragraph()
                    para.style = 'List Bullet'
                    para.add_run(line[2:])
                elif line.startswith('**') and line.endswith('**'):
                    # Bold text
                    para = doc.add_paragraph()
                    run = para.add_run(line[2:-2])
                    run.bold = True
                else:
                    # Regular paragraph
                    doc.add_paragraph(line)
        
        # Save the document
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        clean_title = re.sub(r'[^\w\s-]', '', title).strip()
        clean_title = re.sub(r'[-\s]+', '_', clean_title)
        filename = f"{clean_title}_{timestamp}.docx"
        
        if not os.path.exists(self.downloads_dir):
            os.makedirs(self.downloads_dir)
        
        filepath = os.path.join(self.downloads_dir, filename)
        doc.save(filepath)
        
        return filepath

class CurriculumManager:
    """Manages curriculum context for better relevancy analysis"""
    
    def __init__(self, curriculum_pdf_path: str = None):
        self.curriculum_data = self._extract_curriculum_structure()
    
    def _extract_curriculum_structure(self) -> dict:
        """Extract structured curriculum data from the PDF content"""
        return {
            "modules": {
                "full_stack_llm": {
                    "weeks": "1-6",
                    "description": "UI/APIs, LLMs, prompt engineering, databases, MVP building",
                    "topics": [
                        "UI Building with Python and Gradio",
                        "APIs with FastAPI and CRUD operations", 
                        "LLM & Prompt Engineering fundamentals",
                        "Domain Modeling and Database integration",
                        "AI-accelerated programming with Cursor/Lovable",
                        "GitHub, Open Source, and MVP building"
                    ],
                    "learning_outcomes": [
                        "Set up programming environment for Python",
                        "Build chatbots and apps using Gradio", 
                        "Develop APIs using FastAPI",
                        "Integrate LLMs into applications",
                        "Deploy applications on platforms like Vercel",
                        "Build full-stack AI applications"
                    ]
                },
                "augmented_llm": {
                    "weeks": "7-12", 
                    "description": "Function calling, MCP, LangChain, RAG, fine-tuning",
                    "topics": [
                        "Function/Tool Calling and MCP integration",
                        "LLM Workflows and Chains with LangChain",
                        "Retrieval-Augmented Generation (RAG)",
                        "Advanced RAG with LlamaIndex",
                        "Fine-tuning vs RAG decision frameworks",
                        "GenAI Application Architecture"
                    ],
                    "learning_outcomes": [
                        "Implement function calling in LLM workflows",
                        "Build custom workflows with LangChain",
                        "Master RAG techniques for enhanced LLMs",
                        "Create efficient indexing and retrieval workflows", 
                        "Apply decision frameworks for AI project planning",
                        "Architect scalable GenAI applications"
                    ]
                },
                "ai_agents": {
                    "weeks": "13-16",
                    "description": "Single agents, multi-agent systems, ReAct frameworks", 
                    "topics": [
                        "Introduction to AI Agents and components",
                        "Building AI agents with OpenAI Responses API",
                        "ReAct framework and autonomous decision-making",
                        "Multi-Agent Systems and collaboration protocols",
                        "Agent coordination and communication",
                        "Guardrails, monitoring, and evaluation"
                    ],
                    "learning_outcomes": [
                        "Understand core principles of AI agents",
                        "Build Level-1 agents using Responses API", 
                        "Develop ReAct agents with reasoning capabilities",
                        "Implement multi-agent system architectures",
                        "Apply guardrails for AI safety and security",
                        "Monitor and evaluate agent performance"
                    ]
                },
                "diffusion": {
                    "weeks": "16-17",
                    "description": "SDXL, image generation, ComfyUI workflows",
                    "topics": [
                        "Evolution of GenAI and diffusion model overview",
                        "SDXL prompting and image generation techniques",
                        "ControlNet, inpainting, and guided generation", 
                        "Identity preservation and style transfer",
                        "ComfyUI workflows and automation",
                        "Model deployment on Replicate"
                    ],
                    "learning_outcomes": [
                        "Craft advanced prompts for SDXL",
                        "Apply ControlNet for precise image control",
                        "Train and fine-tune LoRA models",
                        "Build automated ComfyUI workflows",
                        "Deploy production-style demos",
                        "Understand advanced model families"
                    ]
                }
            },
            "student_demographics": {
                "roles": {
                    "engineers": 0.41,
                    "founders": 0.11, 
                    "designers": 0.09,
                    "data_scientists": 0.09,
                    "product_managers": 0.08,
                    "marketing": 0.08,
                    "management": 0.05,
                    "others": 0.09
                },
                "experience_levels": {
                    "0_1_years": 0.20,
                    "1_4_years": 0.30, 
                    "4_10_years": 0.40,
                    "10_plus_years": 0.10
                }
            }
        }
    
    def get_curriculum_context_for_system_prompt(self) -> str:
        """Generate curriculum context string for system prompt enhancement"""
        context = """
## 100xEngineers Curriculum Context

### Module Structure:
1. **Full-Stack LLM (Weeks 1-6)**: UI/APIs, LLMs, prompt engineering, databases, MVP building
   - Key topics: Gradio UIs, FastAPI, LLM integration, domain modeling, AI-accelerated development
   
2. **Augmented LLM (Weeks 7-12)**: Function calling, MCP, LangChain, RAG, fine-tuning  
   - Key topics: Tool integration, LLM workflows, RAG techniques, fine-tuning decisions
   
3. **AI Agents (Weeks 13-16)**: Single agents, multi-agent systems, ReAct frameworks
   - Key topics: Agent components, OpenAI Responses API, ReAct patterns, multi-agent coordination
   
4. **Diffusion (Weeks 16-17)**: SDXL, image generation, ComfyUI workflows
   - Key topics: Image generation, ControlNet, model training, automated workflows

### Student Demographics:
- **Roles**: Engineers (41%), Founders (11%), Designers (9%), Data Scientists (9%), PMs (8%), Marketing (8%), Management (5%), Others (9%)
- **Experience**: 0-1y (20%), 1-4y (30%), 4-10y (40%), 10+y (10%)

Use this context for curriculum relevance scoring in Gate A analysis.
"""
        return context
    
    def map_topic_to_curriculum(self, topic: str) -> dict:
        """Map a topic to relevant curriculum modules and weeks"""
        topic_lower = topic.lower()
        mappings = []
        
        # Check each module for topic relevance
        for module_name, module_data in self.curriculum_data["modules"].items():
            relevance_score = 0
            matched_topics = []
            
            # Check topics within the module
            for curriculum_topic in module_data["topics"]:
                if any(keyword in curriculum_topic.lower() for keyword in topic_lower.split()):
                    relevance_score += 1
                    matched_topics.append(curriculum_topic)
            
            # Check learning outcomes 
            for outcome in module_data["learning_outcomes"]:
                if any(keyword in outcome.lower() for keyword in topic_lower.split()):
                    relevance_score += 0.5
            
            if relevance_score > 0:
                mappings.append({
                    "module": module_name,
                    "weeks": module_data["weeks"], 
                    "description": module_data["description"],
                    "relevance_score": relevance_score,
                    "matched_topics": matched_topics
                })
        
        return {
            "topic": topic,
            "mappings": sorted(mappings, key=lambda x: x["relevance_score"], reverse=True),
            "total_relevance": sum(m["relevance_score"] for m in mappings)
        }

class IntentRecognizer:
    """Recognize user intents from natural language input"""
    
    def __init__(self):
        self.intent_patterns = {
            # Google Drive analysis patterns - very specific
            'analyze': [
                r'brief\s+me\s+on\s+(.+?)\s+(?:from|in|on|based\s+on)\s+(?:my\s+)?(?:google\s+)?drive',
                r'analyze\s+(.+?)\s+(?:from|in|on)\s+(?:my\s+)?(?:google\s+)?drive',
                r'analyze\s+(.+?)\s+(?:documents?|files?)',
                r'analyze\s+my\s+(.+?)\s+(?:documents?|files?)',
                r'study\s+(.+?)\s+(?:from|in)\s+(?:my\s+)?(?:google\s+)?drive',
                r'examine\s+(.+?)\s+(?:from|in)\s+(?:my\s+)?(?:google\s+)?drive',
                r'investigate\s+(.+?)\s+(?:from|in)\s+(?:my\s+)?(?:google\s+)?drive',
                r'can\s+you\s+analyze\s+(.+?)\s+(?:from|in)\s+(?:my\s+)?(?:google\s+)?drive',
                r'please\s+analyze\s+(.+?)\s+(?:from|in)\s+(?:my\s+)?(?:google\s+)?drive'
            ],
            # Web research patterns - for external information
            'research': [
                r'research\s+(.+)',
                r'find\s+(?:the\s+)?latest\s+(?:info(?:rmation)?|news)\s+(?:about|on)\s+(.+)',
                r"what'?s\s+new\s+(?:in|about|with)\s+(.+)",
                r"what'?s\s+the\s+latest\s+(?:on|in|about)\s+(.+)",
                r'latest\s+developments?\s+in\s+(.+)',
                r'current\s+trends\s+in\s+(.+)',
                r'search\s+(?:the\s+web|online)\s+for\s+(.+)',
                r'look\s+up\s+(?:online|current)\s+(.+)'
            ],
            # Knowledge questions - general AI responses
            'question': [
                r'what\s+is\s+(.+?)\s*\??$',
                r'how\s+does\s+(.+?)\s+work\s*\??$',
                r'how\s+to\s+(.+?)\s*\??$',
                r'explain\s+(.+?)(?:\s+to\s+me)?\s*$',
                r'can\s+you\s+explain\s+(.+?)\s*\??$',
                r'tell\s+me\s+about\s+(.+?)(?:\s+please)?\s*$',
                r'describe\s+(.+?)\s*$',
                r'define\s+(.+?)\s*$'
            ],
            # Conversational patterns - general chat
            'conversation': [
                r'^(?:hi|hello|hey)(?:\s+there)?(?:\s+.*)?$',
                r'^(?:good\s+(?:morning|afternoon|evening))(?:\s+.*)?$',
                r'^(?:how\s+are\s+you|how\s+is\s+it\s+going)(?:\s+.*)?$',
                r'^(?:thanks?|thank\s+you)(?:\s+.*)?$',
                r'^(?:bye|goodbye|see\s+you)(?:\s+.*)?$',
                r'^(?:ok|okay|alright|got\s+it|understood)(?:\s+.*)?$',
                r'^(?:help|what\s+can\s+you\s+do)(?:\s+.*)?$'
            ],
            'follow_up': [
                r'tell\s+me\s+more',
                r'elaborate',
                r'expand\s+on\s+that',
                r'more\s+details?',
                r'dive\s+deeper',
                r'what\s+about\s+(.+)',
                r'and\s+what\s+about\s+(.+)',
                r'what\s+are\s+the\s+(.+)',
                r'how\s+about\s+(.+)'
            ],
            'memory_control': [
                r'start\s+(?:a\s+)?new\s+(?:topic|conversation)',
                r'new\s+topic',
                r'clear\s+memory',
                r'forget\s+everything',
                r'end\s+conversation',
                r'start\s+fresh',
                r'remember\s+this'
            ]
        }
    
    def recognize_intent(self, user_input: str) -> dict:
        """Recognize intent and extract topic/entity from user input"""
        user_input = user_input.lower().strip()
        
        for intent, patterns in self.intent_patterns.items():
            for pattern in patterns:
                match = re.search(pattern, user_input, re.IGNORECASE)
                if match:
                    # Extract topic/entity if captured in groups
                    topic = match.group(1).strip() if match.groups() else None
                    
                    return {
                        'intent': intent,
                        'topic': topic,
                        'confidence': 0.9,  # High confidence for pattern match
                        'raw_input': user_input
                    }
        
        # Default intent if no pattern matches
        return {
            'intent': 'general',
            'topic': user_input,
            'confidence': 0.5,
            'raw_input': user_input
        }

class ConversationMemory:
    """Manage conversation context and memory"""
    
    def __init__(self):
        self.current_topic = None
        self.last_analysis = None
        self.conversation_history = []
        self.session_context = {}
    
    def add_interaction(self, user_input: str, response: str, intent: dict):
        """Add an interaction to conversation history"""
        interaction = {
            'timestamp': datetime.now(),
            'user_input': user_input,
            'intent': intent,
            'response': response[:500] + '...' if len(response) > 500 else response  # Truncate long responses
        }
        self.conversation_history.append(interaction)
        
        # Keep only last 10 interactions to manage memory
        if len(self.conversation_history) > 10:
            self.conversation_history = self.conversation_history[-10:]
    
    def set_current_analysis(self, topic: str, analysis: str):
        """Store current analysis for follow-up questions"""
        self.current_topic = topic
        self.last_analysis = analysis
        self.session_context['last_analysis_topic'] = topic
        self.session_context['last_analysis_timestamp'] = datetime.now()
    
    def get_context_for_followup(self) -> str:
        """Generate context string for follow-up questions"""
        if not self.last_analysis:
            return "No previous analysis available for context."
        
        context = f"Previous analysis topic: {self.current_topic}\n"
        context += f"Key points from last analysis:\n{self.last_analysis[:1000]}..."
        return context
    
    def clear_memory(self):
        """Clear conversation memory for new topic"""
        self.current_topic = None
        self.last_analysis = None
        self.conversation_history = []
        self.session_context = {}
    
    def should_start_new_topic(self, new_topic: str) -> bool:
        """Determine if we should start a new conversation topic"""
        if not self.current_topic:
            return True
        
        # Simple topic similarity check
        current_words = set(self.current_topic.lower().split())
        new_words = set(new_topic.lower().split()) 
        overlap = len(current_words.intersection(new_words)) / len(current_words.union(new_words))
        
        return overlap < 0.3  # Start new topic if less than 30% overlap

class ActionDispatcher:
    """Route recognized intents to appropriate actions"""
    
    def __init__(self, config, gdrive_client, report_generator):
        self.config = config
        self.gdrive_client = gdrive_client
        self.report_generator = report_generator
        self.console = Console()
    
    def dispatch_action(self, intent_data: dict, memory: ConversationMemory) -> str:
        """Route intent to appropriate action and return response"""
        intent = intent_data['intent']
        topic = intent_data['topic']
        
        try:
            if intent == 'analyze':
                return self._handle_analyze(topic, memory)
            elif intent == 'research':
                return self._handle_research(topic, memory)
            elif intent == 'question':
                return self._handle_question(topic, memory)
            elif intent == 'conversation':
                return self._handle_conversation(intent_data['raw_input'], memory)
            elif intent == 'follow_up':
                return self._handle_followup(topic, memory)
            elif intent == 'memory_control':
                return self._handle_memory_control(intent_data['raw_input'], memory)
            else:
                return self._handle_general(intent_data['raw_input'], memory)
                
        except Exception as e:
            return f"I encountered an error while processing your request: {str(e)}"
    
    def _handle_analyze(self, topic: str, memory: ConversationMemory) -> str:
        """Handle analysis requests"""
        self.console.print(f"Analyzing {topic} from your Google Drive...")
        
        try:
            # Search Google Drive for relevant documents
            files = self.gdrive_client.list_files(query=topic)
            files = files[:10] if files else []
            
            # Process documents
            documents = []
            for file_info in files:
                if file_info['mimeType'] == 'application/vnd.google-apps.document':
                    text_content = self.gdrive_client.get_file_content(file_info['id'])
                    if text_content:
                        doc_info = DocumentReader.get_document_info(
                            text_content, file_info['mimeType'], file_info['name']
                        )
                        documents.append(doc_info)
            
            if not documents:
                return f"I couldn't find any documents related to '{topic}' in your Google Drive."
            
            # Initialize analyzer with curriculum context
            analyzer = OpenAIAnalyzer(
                self.config.openai_api_key,
                "gpt-4o-mini",
                enable_web_search=True,
                system_prompt_path=self.config.system_prompt_path,
                curriculum_manager=self.config.curriculum_manager
            )
            
            # Perform analysis
            analysis_result = analyzer.analyze_documents_with_web_research(documents, topic)
            
            # Store in memory for follow-ups
            memory.set_current_analysis(topic, analysis_result)
            
            # Generate and save report
            report = self.report_generator.generate_enhanced_report(
                documents, analysis_result, topic, "gpt-4o-mini"
            )
            
            return f"Analysis complete! I've analyzed {len(documents)} documents about '{topic}' and created a comprehensive report. The analysis has been saved to your system and uploaded to Google Drive.\n\nWould you like me to elaborate on any specific aspect of the analysis?"
            
        except Exception as e:
            return f"I encountered an error during analysis: {str(e)}"
    
    def _handle_research(self, topic: str, memory: ConversationMemory) -> str:
        """Handle research requests"""
        self.console.print(f"Researching latest developments in {topic}...")
        
        try:
            analyzer = OpenAIAnalyzer(
                self.config.openai_api_key,
                "gpt-4o-mini", 
                enable_web_search=True,
                system_prompt_path=self.config.system_prompt_path,
                curriculum_manager=self.config.curriculum_manager
            )
            
            result = analyzer.web_research_only(topic)
            
            # Save research result
            kolkata_tz = pytz.timezone('Asia/Kolkata')
            timestamp = datetime.now(kolkata_tz).strftime('%Y-%m-%d %H:%M:%S %Z')
            full_report = f"# Web Research: {topic}\n\nGenerated on: {timestamp}\nModel: gpt-4o-mini with Web Search\n\n{result}"
            
            # Save and upload
            report_gen = ReportGenerator(self.gdrive_client)
            filepath = report_gen._save_report_to_file(full_report, f"Research_{topic}")
            
            if self.gdrive_client:
                self.gdrive_client.create_google_doc_with_links(
                    full_report,
                    f"Web_Research_{topic}_{datetime.now().strftime('%Y%m%d')}"
                )
            
            memory.set_current_analysis(f"Research: {topic}", result)
            
            return f"Research complete! Here are the key findings:\n\n{result}\n\nThe full research report has been saved and uploaded. Would you like me to dive deeper into any particular aspect?"
            
        except Exception as e:
            return f"I encountered an error during research: {str(e)}"
    
    def _handle_question(self, question: str, memory: ConversationMemory) -> str:
        """Handle direct questions using AI knowledge"""
        self.console.print(f"Thinking about: {question}")
        
        try:
            analyzer = OpenAIAnalyzer(
                self.config.openai_api_key,
                self.config.openai_model,
                enable_web_search=False,  # Use AI knowledge, not web search by default
                system_prompt_path=self.config.system_prompt_path,
                curriculum_manager=self.config.curriculum_manager
            )
            
            # Create a knowledge-based question prompt
            question_prompt = f"""Please answer this question using your knowledge: {question}

Provide a helpful and informative response. If the question seems to require analysis of specific documents or real-time information, suggest that the user:
- For document analysis: Ask to "analyze [topic] from my Google Drive" 
- For current information: Ask to "research [topic]" for latest developments

Keep your answer informative but conversational."""
            
            answer = analyzer.generate_response(question_prompt, max_tokens=800)
            
            # Store Q&A in memory
            memory.add_interaction(question, answer, {'intent': 'question'})
            
            return answer
            
        except Exception as e:
            return f"I encountered an issue while thinking about your question: {str(e)}. Could you try rephrasing it?"
    
    def _handle_followup(self, followup: str, memory: ConversationMemory) -> str:
        """Handle follow-up questions using conversation context"""
        if not memory.last_analysis:
            return "I don't have any previous analysis to follow up on. Please ask me to analyze a topic first."
        
        context = memory.get_context_for_followup()
        enhanced_question = f"Based on our previous discussion about {memory.current_topic}: {followup}\n\nContext: {context}"
        
        return self._handle_question(enhanced_question, memory)
    
    def _handle_memory_control(self, command: str, memory: ConversationMemory) -> str:
        """Handle memory control commands"""
        command_lower = command.lower()
        
        if any(phrase in command_lower for phrase in ['new topic', 'start fresh', 'clear memory']):
            memory.clear_memory()
            return "Memory cleared! I'm ready for a new topic. What would you like to explore?"
        elif 'end conversation' in command_lower:
            return "Conversation ended. Thank you for using the AI research assistant!"
        elif 'remember this' in command_lower:
            return "I'll remember our current conversation context for follow-up questions."
        else:
            return "I understand you want to control memory, but I'm not sure what specific action you want."
    
    def _handle_conversation(self, input_text: str, memory: ConversationMemory) -> str:
        """Handle casual conversation"""
        input_lower = input_text.lower().strip()
        
        # Greeting responses
        if any(greeting in input_lower for greeting in ['hi', 'hello', 'hey']):
            return "Hello! I'm your AI workflow co-pilot. I can help you analyze documents from your Google Drive, research topics, or just chat. What would you like to talk about?"
        
        elif any(greeting in input_lower for greeting in ['good morning', 'good afternoon', 'good evening']):
            return "Good day! I'm here to assist you with analysis, research, or conversation. How can I help you today?"
        
        elif any(phrase in input_lower for phrase in ['how are you', 'how is it going']):
            return "I'm doing well, thanks for asking! I'm ready to help you with any document analysis, research, or questions you might have. What's on your mind?"
        
        elif any(thanks in input_lower for thanks in ['thank', 'thanks']):
            return "You're welcome! I'm always here to help. Is there anything else you'd like to discuss or analyze?"
        
        elif any(bye in input_lower for bye in ['bye', 'goodbye', 'see you']):
            return "Goodbye! Feel free to come back anytime you need help with analysis or research. Have a great day!"
        
        elif any(confirm in input_lower for confirm in ['ok', 'okay', 'alright', 'got it', 'understood']):
            return "Great! What would you like to explore next? I can analyze your documents, research current topics, or answer questions."
        
        elif any(help_word in input_lower for help_word in ['help', 'what can you do']):
            return """I'm your AI workflow co-pilot! Here's what I can do:

• **Analyze documents**: "Analyze AI agents from my Google Drive" - I'll search and analyze your documents
• **Research topics**: "Research latest developments in AI" - I'll find current information online  
• **Answer questions**: "What is machine learning?" - I'll explain topics using my knowledge
• **Chat naturally**: Just talk to me like you would with any assistant

Simply speak naturally - I'll understand your intent and respond accordingly!"""
        
        else:
            return "I understand you're trying to tell me something, but I'm not quite sure what you need. You can ask me to analyze documents from your Google Drive, research topics online, ask questions, or just chat. What would you like to do?"

    def _handle_general(self, input_text: str, memory: ConversationMemory) -> str:
        """Handle inputs that don't match specific intents - use AI for general conversation"""
        # Instead of defaulting to analysis, use AI to have a natural conversation
        try:
            analyzer = OpenAIAnalyzer(
                self.config.openai_api_key,
                self.config.openai_model,
                enable_web_search=False,  # Don't use web search for general conversation
                system_prompt_path=self.config.system_prompt_path,
                curriculum_manager=None  # No curriculum context for simple conversations
            )
            
            # Create a conversational prompt
            conversation_prompt = f"""You are a helpful AI workflow co-pilot assistant. The user said: "{input_text}"

Please respond naturally and conversationally. If they seem to be asking for document analysis from Google Drive, suggest they be more specific (e.g., "analyze X from my Google Drive"). If they want research on current topics, suggest they ask for research on specific topics.

Keep responses friendly and helpful, but don't automatically start analyzing documents or searching unless explicitly requested with clear intent."""

            response = analyzer.generate_response(conversation_prompt, max_tokens=500)
            return response
            
        except Exception as e:
            return "I'm not sure what you'd like me to do. You can ask me to analyze topics from your Google Drive, research current subjects, ask questions, or just chat. What interests you today?"

class ConversationManager:
    """Main orchestrator for conversational interactions"""
    
    def __init__(self, config):
        self.config = config
        self.intent_recognizer = IntentRecognizer()
        self.memory = ConversationMemory()
        self.console = Console()
        
        # Initialize clients
        try:
            self.gdrive_client = GoogleDriveClient(
                credentials_path=config.google_credentials_path,
                drive_token_path=config.google_drive_token_path,
                docs_token_path=config.google_docs_token_path
            )
            self.report_generator = ReportGenerator(self.gdrive_client)
            self.action_dispatcher = ActionDispatcher(config, self.gdrive_client, self.report_generator)
        except Exception as e:
            self.console.print(f"Warning: Could not initialize Google Drive client: {e}")
            self.gdrive_client = None
            self.report_generator = ReportGenerator()
            self.action_dispatcher = ActionDispatcher(config, None, self.report_generator)
    
    def start_conversation(self):
        """Start the main conversation loop"""
        self.console.print("\n" + "="*60)
        self.console.print("AI Research Assistant - Second Brain Mode", style="bold blue")
        self.console.print("="*60)
        self.console.print("\nHello! I'm your AI research assistant with access to your Google Drive and curriculum context.")
        self.console.print("\nYou can:")
        self.console.print("• Analyze topics from your Google Drive")  
        self.console.print("• Research latest developments")
        self.console.print("• Ask questions about your documents")
        self.console.print("• Have follow-up conversations")
        self.console.print("\nJust type naturally - no commands needed!")
        self.console.print("Type 'exit' or 'quit' to end the conversation.\n")
        
        while True:
            try:
                # Get user input
                user_input = input("\n> ").strip()
                
                # Check for exit commands
                if user_input.lower() in ['exit', 'quit', 'bye', 'goodbye']:
                    self.console.print("\nThank you for using the AI Research Assistant! Goodbye!")
                    break
                
                if user_input.lower() in ['end', 'end conversation']:
                    self.console.print("\nConversation ended. Ready for the next topic!")
                    self.memory.conversations = []  # Clear conversation history
                    continue
                
                if not user_input:
                    continue
                
                # Recognize intent
                intent_data = self.intent_recognizer.recognize_intent(user_input)
                
                # Dispatch action and get response
                response = self.action_dispatcher.dispatch_action(intent_data, self.memory)
                
                # Display response
                self.console.print(f"\n{response}")
                
                # Store interaction in memory
                self.memory.add_interaction(user_input, response, intent_data)
                
            except KeyboardInterrupt:
                self.console.print("\n\nConversation interrupted. Goodbye!")
                break
            except Exception as e:
                self.console.print(f"\nError: {str(e)}")
                continue

# =============================================================================
# VOICE-READY ARCHITECTURE (Step 5)
# =============================================================================

class InputInterface:
    """Abstract interface for user input"""
    def get_user_input(self) -> str:
        raise NotImplementedError

class TextInput(InputInterface):
    """Text-based input implementation"""
    def get_user_input(self) -> str:
        return input("\n> ").strip()

class VoiceInput(InputInterface):
    """Voice input implementation using LiveKit STT"""
    def __init__(self):
        self.stt = None
        if LIVEKIT_AVAILABLE:
            try:
                self.stt = lk_openai.STT(model="gpt-4o-transcribe")
            except Exception as e:
                print(f"Warning: Could not initialize STT: {e}")
    
    def get_user_input(self) -> str:
        if self.stt and LIVEKIT_AVAILABLE:
            # LiveKit STT implementation would go here
            # For now, fallback to text input with voice indicator
            return input("\n[Voice] > ").strip()
        else:
            return input("\n[Voice] > ").strip()

class OutputInterface:
    """Abstract interface for response output"""
    def send_response(self, text: str):
        raise NotImplementedError

class TextOutput(OutputInterface):
    """Text-based output implementation"""
    def __init__(self):
        self.console = Console()
    
    def send_response(self, text: str):
        self.console.print(f"\n{text}")

class VoiceOutput(OutputInterface):
    """Voice output implementation using LiveKit TTS"""
    def __init__(self):
        self.console = Console()
        self.tts = None
        if LIVEKIT_AVAILABLE:
            try:
                self.tts = lk_openai.TTS(
                    model="gpt-4o-mini-tts",
                    voice="ash",
                    instructions="Speak quickly and clearly; be concise and confident.",
                    speed=1.2,
                )
            except Exception as e:
                print(f"Warning: Could not initialize TTS: {e}")
    
    def send_response(self, text: str):
        # Display text
        self.console.print(f"\n[Audio] {text}")
        
        if self.tts and LIVEKIT_AVAILABLE:
            # LiveKit TTS implementation would go here
            # For now, just display with voice indicator
            pass

# =============================================================================
# MAIN CONVERSATIONAL AI SYSTEM (Step 6 Integration)
# =============================================================================

class ConversationalGDriveAnalyzer:
    """Main conversational AI system integrating all components"""
    
    def __init__(self, voice_mode=False):
        self.config = Config()
        self.conversation_manager = ConversationManager(self.config)  # Use simple version
        self.voice_mode = voice_mode and LIVEKIT_AVAILABLE
        
        # Voice-ready interfaces
        if self.voice_mode:
            self.input_interface = VoiceInput()
            self.output_interface = VoiceOutput()
        else:
            self.input_interface = TextInput()
            self.output_interface = TextOutput()
    
    def run(self):
        """Main entry point for conversational mode"""
        try:
            if self.voice_mode:
                print("Voice mode enabled (LiveKit integration)")
            self.conversation_manager.start_conversation()
        except Exception as e:
            print(f"Error starting conversation: {e}")

# =============================================================================
# LIVEKIT VOICE INTEGRATION
# =============================================================================
"""
async def livekit_entrypoint(ctx):
    # LiveKit entry point for voice-based conversations
    if not LIVEKIT_AVAILABLE:
        raise ImportError("LiveKit not available. Install with: pip install livekit-agents")
    
    await ctx.connect(auto_subscribe="AUDIO_ONLY")
    
    # Create a voice-enabled agent session
    session = AgentSession(
        stt=lk_openai.STT(model="gpt-4o-transcribe"),
        llm=lk_openai.LLM(model="gpt-4o-mini"),
        tts=lk_openai.TTS(
            model="gpt-4o-mini-tts",
            voice="ash",
            instructions="Speak quickly and clearly; be concise and confident.",
            speed=1.2,
        ),
        vad=silero.VAD.load(),
        turn_detection="vad",
    )"""
    
   

# =============================================================================
# PHASE 2: PATTERN LEARNING ENGINE
# =============================================================================

import json
from pathlib import Path

class PatternLearningEngine:
    """Learn patterns from test cases to build AI reasoning capabilities"""
    
    def __init__(self):
        self.learned_patterns = {
            'thought_processes': [],
            'relevancy_mappings': [],
            'risk_assessments': [], 
            'hypothesis_structures': [],
            'evidence_analysis_patterns': []
        }
        self.pattern_file = Path("learned_patterns.json")
        self.load_patterns()
    
    def load_patterns(self):
        """Load previously learned patterns"""
        if self.pattern_file.exists():
            try:
                with open(self.pattern_file, 'r') as f:
                    self.learned_patterns = json.load(f)
            except Exception as e:
                print(f"Could not load patterns: {e}")
    
    def save_patterns(self):
        """Save learned patterns to file"""
        try:
            with open(self.pattern_file, 'w') as f:
                json.dump(self.learned_patterns, f, indent=2)
        except Exception as e:
            print(f"Could not save patterns: {e}")
    
    def learn_from_test_case(self, input_problem: str, thought_process: str, expected_output: str):
        """Learn patterns from a single test case"""
        # Extract thought process patterns
        if "First principles:" in thought_process:
            pattern = self._extract_first_principles_pattern(thought_process)
            self.learned_patterns['thought_processes'].append(pattern)
        
        # Extract hypothesis patterns
        if "Output/hypothesis" in expected_output or "Treat" in expected_output:
            pattern = self._extract_hypothesis_pattern(expected_output)
            self.learned_patterns['hypothesis_structures'].append(pattern)
        
        # Extract risk assessment patterns
        if "Risks/unknowns" in thought_process.lower():
            pattern = self._extract_risk_pattern(thought_process)
            self.learned_patterns['risk_assessments'].append(pattern)
        
        self.save_patterns()
    
    def _extract_first_principles_pattern(self, thought_process: str) -> dict:
        """Extract first principles reasoning pattern"""
        return {
            'type': 'first_principles',
            'indicators': ['First principles:', 'Translation:', 'enablers', 'adoption'],
            'structure': 'problem_statement -> enablers -> translation -> adoption_pattern',
            'confidence': 0.8
        }
    
    def _extract_hypothesis_pattern(self, output: str) -> dict:
        """Extract hypothesis formation pattern"""
        return {
            'type': 'hypothesis',
            'structure': 'reframe -> actionable_recommendation -> time_bound_prediction',
            'indicators': ['Treat', 'Build', 'Bet:', 'within', 'months'],
            'confidence': 0.9
        }
    
    def _extract_risk_pattern(self, thought_process: str) -> dict:
        """Extract risk assessment pattern"""
        return {
            'type': 'risk_assessment', 
            'categories': ['source_quality', 'cost_time', 'citations', 'trust_building'],
            'structure': 'identify -> mitigate -> trust_measures',
            'confidence': 0.7
        }
    
    def apply_patterns_to_analysis(self, topic: str, base_analysis: str) -> str:
        """Apply learned patterns to enhance analysis"""
        enhanced_analysis = base_analysis
        
        # Apply thought process patterns
        if self.learned_patterns['thought_processes']:
            enhanced_analysis = self._apply_first_principles_pattern(enhanced_analysis, topic)
        
        # Apply hypothesis patterns
        if self.learned_patterns['hypothesis_structures']:
            enhanced_analysis = self._apply_hypothesis_pattern(enhanced_analysis, topic)
        
        return enhanced_analysis
    
    def _apply_first_principles_pattern(self, analysis: str, topic: str) -> str:
        """Apply first principles thinking pattern"""
        pattern_prompt = f"""
        Apply first principles analysis to {topic}:
        1. Problem-of-the-world: What core constraint existed before this?
        2. Feasible-now enablers: What shifts unlocked viability?
        3. Translation: Simplify to core pattern
        4. Adoption catalysts: What accelerated uptake?
        
        {analysis}
        """
        return pattern_prompt
    
    def _apply_hypothesis_pattern(self, analysis: str, topic: str) -> str:
        """Apply learned hypothesis formation pattern"""
        hypothesis_prompt = f"""
        Generate hypothesis using learned patterns:
        - Reframe: Treat {topic} as [reference implementation/pattern/standard]
        - Actionable bet: [specific recommendation with timeframe]
        - Market prediction: [2-6 month outlook]
        
        {analysis}
        """
        return hypothesis_prompt

class TestCaseProcessor:
    """Process and learn from test case examples"""
    
    def __init__(self, pattern_engine: PatternLearningEngine):
        self.pattern_engine = pattern_engine
        self.test_cases = []
    
    def add_test_case(self, case_number: int, input_problem: str, thought_process: str, expected_output: str):
        """Add a test case for learning"""
        test_case = {
            'number': case_number,
            'input': input_problem,
            'thought_process': thought_process,
            'expected_output': expected_output,
            'learned': False
        }
        self.test_cases.append(test_case)
        
        # Learn from this test case immediately
        self.pattern_engine.learn_from_test_case(input_problem, thought_process, expected_output)
        test_case['learned'] = True
        
        return f"Learned from Test Case {case_number}. Pattern extracted and saved."
    
    def process_example_cases(self):
        """Process test cases from TEST CASES_learning.docx file"""
        test_cases_file = "TEST CASES_learning.docx"
        
        if not os.path.exists(test_cases_file):
            print(f"Warning: {test_cases_file} not found. Using fallback examples.")
            self._load_fallback_cases()
            return
        
        try:
            from docx import Document
            doc = Document(test_cases_file)
            
            current_case_number = None
            current_input = ""
            current_thought = ""
            current_output = ""
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                # Detect test case header
                if text.startswith("Test Case") and "–" in text:
                    # Save previous case if complete
                    if current_case_number and current_input and current_thought and current_output:
                        self.add_test_case(current_case_number, current_input, current_thought, current_output)
                    
                    # Extract case number
                    try:
                        current_case_number = int(text.split()[2])
                    except:
                        current_case_number = len(self.test_cases) + 1
                    
                    current_input = ""
                    current_thought = ""
                    current_output = ""
                
                elif text.startswith("Input/problem"):
                    current_input = text.replace("Input/problem.", "").strip()
                
                elif text.startswith("Thought process"):
                    current_thought = text.replace("Thought process (teach the agent).", "").strip()
                
                elif text.startswith("Output/hypothesis"):
                    current_output = text.replace("Output/hypothesis.", "").strip()
                
                else:
                    # Continue adding to the current section
                    if current_input and not current_thought and not current_output:
                        current_input += " " + text
                    elif current_thought and not current_output:
                        current_thought += " " + text
                    elif current_output:
                        current_output += " " + text
            
            # Add the final case
            if current_case_number and current_input and current_thought and current_output:
                self.add_test_case(current_case_number, current_input, current_thought, current_output)
            
            # Silently load test cases
            
        except Exception as e:
            print(f"Error reading {test_cases_file}: {e}")
            self._load_fallback_cases()
    
    def _load_fallback_cases(self):
        """Load fallback cases if file reading fails"""
        self.add_test_case(
            1,
            "Is 'deep research' a product we should copy or a capability to integrate?",
            """Deep Research is OpenAI's agentic mode that decomposes a query, browses, and synthesizes a source-rich report; doc guidance stresses specificity, constraints, and evaluation; updates add visual browsing and API access; o-series reasoning models (o3/o4-mini) drive longer deliberation. Translation: "planning + browsing + synthesis" is a pattern you can reproduce, not just a SKU to rent.""",
            """Treat Deep Research as a reference implementation. Build your own agent with explicit planning, source selection, and synthesis suited to your curriculum and evals."""
        )

class LearningConversationManager(ConversationManager):
    """Enhanced ConversationManager with pattern learning capabilities"""
    
    def __init__(self, config):
        super().__init__(config)
        self.pattern_engine = PatternLearningEngine()
        self.test_case_processor = TestCaseProcessor(self.pattern_engine)
        self.learning_mode = False
    
    def start_conversation(self):
        """Start conversation with learning capabilities"""
        self.console.print("\n" + "="*70)
        self.console.print("AI COPILOT - Second Brain Learning Mode", style="bold blue")
        self.console.print("="*70)
        self.console.print("\nHello! I'm your AI research COPILOT that learns from your expertise.")
        self.console.print("\nCapabilities:")
        self.console.print("• Analyze topics from Google Drive with curriculum context")
        self.console.print("• Research latest developments with balanced historical data")
        self.console.print("• Learn from your test cases to improve reasoning")
        self.console.print("• Apply learned patterns to generate better analyses")
        self.console.print("\nLearning Commands:")
        self.console.print("• 'teach me test case X: [input] -> [thought] -> [output]'")
        self.console.print("• 'load example test cases' - Process your provided examples")
        self.console.print("• 'apply learned patterns to [topic]' - Use learned reasoning")
        self.console.print("\nJust type naturally - I'll understand!")
        self.console.print("Type 'exit' or 'quit' to end.\n")
        
        # Load test cases silently
        self.test_case_processor.process_example_cases()
        
        while True:
            try:
                user_input = input("\n> ").strip()
                
                if user_input.lower() in ['exit', 'quit', 'bye', 'goodbye']:
                    self.console.print("\nThank you for using the Learning AI Assistant! Goodbye!")
                    break
                
                if user_input.lower() in ['end', 'end conversation']:
                    self.console.print("\nConversation ended. Ready for the next topic!")
                    self.memory.conversations = []  # Clear conversation history
                    continue
                
                if not user_input:
                    continue
                
                # Check for learning commands first
                if self._handle_learning_commands(user_input):
                    continue
                
                # Regular conversation flow with pattern enhancement
                intent_data = self.intent_recognizer.recognize_intent(user_input)
                response = self.action_dispatcher.dispatch_action(intent_data, self.memory)
                
                # Apply learned patterns to enhance response
                if hasattr(self, 'pattern_engine') and intent_data.get('topic'):
                    response = self.pattern_engine.apply_patterns_to_analysis(intent_data['topic'], response)
                
                self.console.print(f"\n{response}")
                self.memory.add_interaction(user_input, response, intent_data)
                
            except KeyboardInterrupt:
                self.console.print("\n\nConversation interrupted. Goodbye!")
                break
            except Exception as e:
                self.console.print(f"\nError: {str(e)}")
                continue
    
    def _handle_learning_commands(self, user_input: str) -> bool:
        """Handle learning-specific commands"""
        user_lower = user_input.lower()
        
        if 'load example test cases' in user_lower:
            self.test_case_processor.process_example_cases()
            return True
        
        if 'teach me test case' in user_lower:
            self.console.print("Please provide test case in format: 'Test Case X: [input] -> [thought] -> [output]'")
            return True
        
        if 'apply learned patterns' in user_lower:
            topic = user_input.replace('apply learned patterns to', '').strip()
            if topic:
                enhanced = self.pattern_engine.apply_patterns_to_analysis(topic, f"Analyzing {topic}")
                self.console.print(f"Applied learned patterns to analyze {topic}:\n{enhanced}")
            return True
        
        return False

class Config:
    """Configuration management for the application"""
    
    def __init__(self):
        load_dotenv()
        self.google_credentials_path = os.getenv("GOOGLE_CREDENTIALS_PATH", "credentials.json")
        self.google_docs_token_path = os.getenv("GOOGLE_DOCS_TOKEN", "docstoken.json")
        self.google_drive_token_path = os.getenv("GOOGLE_DRIVE_TOKEN", "token.json")
        self.openai_api_key = os.getenv("OPENAI_API_KEY")
        self.openai_model = os.getenv("OPENAI_MODEL", "gpt-4o-mini")
        self.system_prompt_path = os.getenv("SYSTEM_PROMPT_PATH")
        self.curriculum_manager = CurriculumManager()
        self._validate_config()
    
    def _validate_config(self):
        """Validate configuration settings"""
        if not os.path.exists(self.google_credentials_path):
            raise FileNotFoundError(
                f"Google credentials file not found at {self.google_credentials_path}. "
                "Please download credentials.json from Google Cloud Console."
            )
        
        if not self.openai_api_key:
            raise ValueError(
                "OpenAI API key not found. Please set OPENAI_API_KEY in your .env file."
            )

# --- Voice Agent (top-level) ---
from livekit.agents import Agent, RoomInputOptions, function_tool, RunContext
from datetime import datetime

class VoiceGDriveAgent(Agent):
    def __init__(self):
        super().__init__(instructions=(
            "You are an AI research assistant with access to Google Drive documents. "
            "You can analyze topics, research developments, ask questions about documents, and generate reports. "
            "Keep responses concise and clear for voice interaction."
        ))
        # Lazy init
        self.config = None
        self.gdrive_client = None
        self.analyzer = None

        # Optional: pattern engine (safe if present)
        try:
            self.pattern_engine = PatternLearningEngine()
            try:
                test_processor = TestCaseProcessor(self.pattern_engine)
                test_processor.process_example_cases()
            except Exception:
                pass
        except Exception:
            self.pattern_engine = None

    @function_tool(
        description=(
            "Analyze a topic using Google Drive docs and create a Google Doc report. "
            "Use when the user says things like 'brief me on X from my drive'."
        )
    )
    async def analyze_drive_topic(
        self, context: RunContext, topic: str, query: str = "", docs_only: bool = False
    ) -> str:
        try:
            # lazy init
            if self.config is None:
                self.config = Config()

            if self.gdrive_client is None:
                try:
                    self.gdrive_client = GoogleDriveClient(
                        credentials_path=self.config.google_credentials_path,
                        drive_token_path=self.config.google_drive_token_path,
                        docs_token_path=self.config.google_docs_token_path
                    )
                except Exception as e:
                    if docs_only:
                        return f"Google Drive initialization failed: {e}"
                    self.gdrive_client = None

            if self.analyzer is None:
                try:
                    self.analyzer = OpenAIAnalyzer(
                        self.config.openai_api_key,
                        self.config.openai_model,
                        system_prompt_path=self.config.system_prompt_path,
                        curriculum_manager=self.config.curriculum_manager
                    )
                except Exception as e:
                    return f"OpenAI analyzer initialization failed: {e}"

            # ---- fetch candidate docs ----
            documents = []
            if self.gdrive_client:
                q = (query or topic).strip()
                try:
                    files = self.gdrive_client.list_files(query=q) or []
                except Exception:
                    files = []
                files = files[:10]  # keep it light for voice

                for f in files:
                    try:
                        if f.get("mimeType") == "application/vnd.google-apps.document":
                            text = self.gdrive_client.get_file_content(f["id"])
                            if text:
                                documents.append(
                                    DocumentReader.get_document_info(text, f["mimeType"], f["name"])
                                )
                        else:
                            content = self.gdrive_client.download_file(f["id"], f["name"])
                            if content:
                                documents.append(
                                    DocumentReader.get_document_info(content, f["mimeType"], f["name"])
                                )
                    except Exception:
                        continue

            if not documents and docs_only:
                return f"No Drive documents found for '{topic}'."

            # ---- run analysis (FIXED SIGNATURE) ----
            if docs_only:
                analysis_text = self.analyzer.analyze_documents(documents, topic)
            else:
                analysis_text = self.analyzer.analyze_documents_with_web_research(documents, topic)

            # ---- upload Google Doc report ----
            gdoc_title = f"{topic} — Drive Brief — {datetime.now().strftime('%Y-%m-%d %H:%M')}"
            uploaded = None
            if self.gdrive_client:
                try:
                    uploaded = self.gdrive_client.upload_document_to_drive(analysis_text, gdoc_title)
                except Exception as e:
                    return f"(Upload failed: {e})\n\n{analysis_text}"

            if uploaded and isinstance(uploaded, dict):
                return f"Created report: {uploaded.get('webViewLink', '(no link)')}"

            return analysis_text

        except Exception as e:
            return f"Voice tool error: {e}"


# =============================================================================
# LIVEKIT VOICE INTEGRATION
# =============================================================================

async def livekit_entrypoint(ctx):
    """Properly configured voice-enabled entrypoint"""
    try:
        print("[ENTRYPOINT] Starting voice agent...")
        
        @ctx.room.on("track_subscribed")
        def on_track_subscribed(track, publication, participant):
            print(f"[DEBUG] Track subscribed: {track.kind} from {participant.identity}")
            if track.kind == "audio":
                print("[DEBUG] Audio track received!")
        
        @ctx.room.on("track_published")  
        def on_track_published(publication, participant):
            print(f"[DEBUG] Track published: {publication.kind} by {participant.identity}")
        
        await ctx.connect(auto_subscribe="audio_only")
        print("[ENTRYPOINT] Connected to room")
        
        # Import the proper voice pipeline
        from livekit.agents import VoicePipelineAgent
        
        # Initialize voice components with proper configuration
        vad = silero.VAD.load()
        stt = lk_openai.STT(model="whisper-1")  # or "gpt-4o-audio-preview" if available
        llm = lk_openai.LLM(model="gpt-4o-mini")
        tts = lk_openai.TTS(
            model="tts-1",  # or "tts-1-hd" for better quality
            voice="alloy",  # or "echo", "fable", "onyx", "nova", "shimmer"
            speed=1.0
        )
        
        # Create initial chat context
        initial_ctx = lk_openai.ChatContext().append(
            role="system",
            text=(
                "You are a helpful AI assistant that can analyze Google Drive documents. "
                "Keep responses concise and clear for voice interaction. "
                "When asked to analyze documents, search the user's Google Drive and provide insights."
            )
        )
        
        # Create the voice pipeline agent
        assistant = VoicePipelineAgent(
            vad=vad,
            stt=stt,
            llm=llm,
            tts=tts,
            chat_ctx=initial_ctx,
            allow_interruptions=True,
            interrupt_speech_duration=0.5,
            interrupt_min_words=3,
        )
        
        # Start the assistant
        assistant.start(ctx.room)
        
        # Initial greeting via voice
        await assistant.say(
            "Hello! I'm ready to help you analyze your Google Drive documents. "
            "Just tell me what you'd like to explore.",
            allow_interruptions=True
        )
        
        print("[ENTRYPOINT] Voice assistant started successfully")
        
    except Exception as e:
        print(f"[ENTRYPOINT] Error: {e}")
        import traceback
        traceback.print_exc()
# =============================================================================
# GOOGLE DRIVE CLIENT
# =============================================================================

class GoogleDriveClient:
    """Handles Google Drive API interactions"""
    
    DRIVE_SCOPES = [
        'https://www.googleapis.com/auth/drive.readonly',
        'https://www.googleapis.com/auth/drive.file'
    ]
    
    DOCS_SCOPES = [
        'https://www.googleapis.com/auth/documents'
    ]
    
    def __init__(self, credentials_path: str = 'credentials.json', 
                 drive_token_path: str = 'token.json',
                 docs_token_path: str = 'docstoken.json'):
        self.credentials_path = credentials_path
        self.drive_token_path = drive_token_path
        self.docs_token_path = docs_token_path
        self.service = None

        self.docs_service = None
        self._authenticate()
    
    def _authenticate(self):
        """Authenticate with Google Drive and Docs APIs separately"""
        # Authenticate Google Drive
        drive_creds = self._get_credentials(self.drive_token_path, self.DRIVE_SCOPES, "Google Drive")
        self.service = self._make_service('drive', 'v3', drive_creds, timeout=180)
        
        # Authenticate Google Docs  
        docs_creds = self._get_credentials(self.docs_token_path, self.DOCS_SCOPES, "Google Docs")
        self.docs_service = self._make_service('docs', 'v1', docs_creds, timeout=180)
    
    def _make_service(self, api: str, version: str, creds, timeout: int = 180):
        """Create service with custom HTTP timeout"""
        http = httplib2.Http(timeout=timeout)
        authed = AuthorizedHttp(creds, http=http)
        return build(api, version, http=authed, cache_discovery=False)
    
    def _get_credentials(self, token_path: str, scopes: list, service_name: str):
        """Get credentials for a specific service"""
        creds = None
        
        
        # Load existing token if it exists
        if os.path.exists(token_path):
            try:
                creds = Credentials.from_authorized_user_file(token_path, scopes)
            except Exception as e:
                logger.warning(f"Failed to load existing {service_name} token: {e}")
        
        # Check if credentials are valid
        if not creds or not creds.valid:
            if creds and creds.expired and creds.refresh_token:
                try:
                    creds.refresh(Request())
                    logger.info(f"Refreshed {service_name} token successfully")
                except Exception as e:
                    logger.warning(f"Failed to refresh {service_name} token: {e}")
                    creds = None
            
            # If we still don't have valid credentials, authenticate
            if not creds or not creds.valid:
                logger.info(f"Authenticating {service_name}...")
                flow = InstalledAppFlow.from_client_secrets_file(self.credentials_path, scopes)
                
                # Try local server first, fallback to console-based auth
                try:
                    creds = flow.run_local_server(port=0)
                except Exception as e:
                    logger.warning(f"Local server authentication failed for {service_name}: {e}")
                    logger.info(f"Falling back to console-based authentication for {service_name}...")
                    creds = flow.run_console()
            
            # Save the credentials
            try:
                # Create directory if it doesn't exist
                token_dir = os.path.dirname(os.path.abspath(token_path))
                if token_dir and not os.path.exists(token_dir):
                    os.makedirs(token_dir)
                
                with open(token_path, 'w') as token:
                    token.write(creds.to_json())
                logger.info(f"Saved {service_name} token to {token_path}")
            except Exception as e:
                logger.error(f"Failed to save {service_name} token: {e}")
        
        return creds
    
    def list_files(
        self,
        query: Optional[str] = None,
        file_types: Optional[List[str]] = None,
        page_size: int = 50,
    ) -> List[Dict]:
        """Search Google Drive. Supports raw Drive q or plain keywords."""
        if file_types is None:
            file_types = [
                "application/pdf",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "application/msword",
                "text/plain",
                "application/vnd.google-apps.document",
            ]

        if not self.service:
            raise RuntimeError("Drive service not initialized")

        # Build Drive q
        parts = ["trashed = false"]  # exclude trash by default
        mime_q = " or ".join([f"mimeType='{mt}'" for mt in file_types])
        parts.append(f"({mime_q})")

        def looks_like_drive_q(s: str) -> bool:
            s = s.lower()
            return any(k in s for k in (" fulltext ", " name ", " mimetype ", " modifiedtime ", " parents ", " trashed ", " in ", " and ", " or "))

        if query:
            q = query.strip()
            if looks_like_drive_q(q):
                parts.append(f"({q})")
            else:
                import re
                tokens = re.findall(r'"[^"]+"|\S+', q)
                terms = []
                for t in tokens:
                    kw = t.strip().strip('"').strip("'")
                    if len(kw) <= 2:
                        continue
                    kw = kw.replace("'", "\\'")
                    terms.append(f"(name contains '{kw}' or fullText contains '{kw}')")
                if terms:
                    parts.append("(" + " and ".join(terms) + ")")

        q_final = " and ".join(parts)

        fields = "nextPageToken, files(id, name, mimeType, size, modifiedTime, webViewLink)"
        results = []
        next_token = None

        while True:
            resp = self.service.files().list(
                q=q_final,
                spaces="drive",
                pageSize=min(page_size, 1000),
                fields=fields,
                orderBy="modifiedTime desc",
                includeItemsFromAllDrives=True,
                supportsAllDrives=True,
                pageToken=next_token,
            ).execute()
            results.extend(resp.get("files", []))
            next_token = resp.get("nextPageToken")
            if not next_token or len(results) >= page_size:
                break

        return results[:page_size]
    
    def download_file(self, file_id: str, file_name: str) -> Optional[bytes]:
        """Download file content from Google Drive with enhanced retry logic"""
        max_retries = 5
        retry_delay = 2  # Base delay in seconds
        
        for attempt in range(max_retries):
            try:
                file_metadata = self.service.files().get(fileId=file_id).execute()
                mime_type = file_metadata.get('mimeType')
                
                if mime_type == 'application/vnd.google-apps.document':
                    request = self.service.files().export_media(
                        fileId=file_id,
                        mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                    )
                else:
                    request = self.service.files().get_media(fileId=file_id)
                
                file_io = io.BytesIO()
                # Use 5MB chunks and add chunk-level retries
                downloader = MediaIoBaseDownload(file_io, request, chunksize=5 * 1024 * 1024)
                
                done = False
                while done is False:
                    status, done = downloader.next_chunk(num_retries=5)
                
                logger.info(f"Successfully downloaded: {file_name}")
                return file_io.getvalue()
                
            except Exception as e:
                logger.warning(f"Attempt {attempt + 1} failed for {file_name}: {str(e)}")
                if attempt < max_retries - 1:
                    # Exponential backoff: 2s, 4s, 8s, 16s, 32s
                    backoff_time = retry_delay * (2 ** attempt)
                    logger.info(f"Retrying {file_name} in {backoff_time}s...")
                    time.sleep(backoff_time)
                else:
                    logger.error(f"Failed to download {file_name} after {max_retries} attempts")
                    return None
        
        return None
    
    def get_file_content(self, file_id: str) -> Optional[str]:
        """Get rich text content for Google Docs using Docs API with comprehensive retry logic"""
        try:
            file_metadata = self.service.files().get(fileId=file_id).execute()
            mime_type = file_metadata.get('mimeType')
            
            if mime_type == 'application/vnd.google-apps.document':
                # Try Docs API with retries
                for attempt in range(5):
                    try:
                        logger.debug(f"Docs API attempt {attempt+1} for {file_id}")
                        document = self.docs_service.documents().get(documentId=file_id).execute()
                        content = self._extract_text_from_doc(document)
                        logger.info(f"Successfully read Google Doc via Docs API: {file_id}")
                        return content
                    except Exception as docs_error:
                        logger.warning(f"Docs API attempt {attempt+1} failed for {file_id}: {docs_error}")
                        if attempt == 4:
                            logger.warning(f"Docs API failed for {file_id}, falling back to export")
                            break
                        else:
                            time.sleep(2 * (2 ** attempt))  # Exponential backoff
                            continue
                
                # Fallback export with robust download
                logger.info(f"Using export fallback for {file_id}")
                for attempt in range(5):
                    try:
                        request = self.service.files().export_media(
                            fileId=file_id,
                            mimeType='text/plain'
                        )
                        file_io = io.BytesIO()
                        downloader = MediaIoBaseDownload(file_io, request, chunksize=5 * 1024 * 1024)
                        
                        done = False
                        while done is False:
                            status, done = downloader.next_chunk(num_retries=5)
                        
                        return file_io.getvalue().decode('utf-8', errors='replace')
                    except Exception as export_error:
                        logger.warning(f"Export attempt {attempt+1} failed for {file_id}: {export_error}")
                        if attempt < 4:
                            time.sleep(2 * (2 ** attempt))
                        else:
                            raise export_error
            
            elif mime_type == 'text/plain':
                for attempt in range(5):
                    try:
                        request = self.service.files().get_media(fileId=file_id)
                        file_io = io.BytesIO()
                        downloader = MediaIoBaseDownload(file_io, request, chunksize=5 * 1024 * 1024)
                        
                        done = False
                        while done is False:
                            status, done = downloader.next_chunk(num_retries=5)
                        
                        return file_io.getvalue().decode('utf-8', errors='replace')
                    except Exception as e:
                        logger.warning(f"Text file attempt {attempt+1} failed for {file_id}: {e}")
                        if attempt < 4:
                            time.sleep(2 * (2 ** attempt))
                        else:
                            raise e
            
            else:
                logger.debug(f"Unsupported mime type for text extraction: {mime_type}")
                return None
                
        except Exception as e:
            logger.error(f"Error getting file content for {file_id}: {str(e)}")
            return None
    
    def _extract_text_from_doc(self, document: dict) -> str:
        """Extract text from Google Docs document structure"""
        content = document.get('body', {}).get('content', [])
        text_parts = []
        
        for element in content:
            if 'paragraph' in element:
                paragraph = element['paragraph']
                elements = paragraph.get('elements', [])
                
                paragraph_text = []
                for elem in elements:
                    if 'textRun' in elem:
                        text_content = elem['textRun'].get('content', '')
                        paragraph_text.append(text_content)
                
                text_parts.append(''.join(paragraph_text))
            
            elif 'table' in element:
                table = element['table']
                for row in table.get('tableRows', []):
                    for cell in row.get('tableCells', []):
                        cell_content = cell.get('content', [])
                        cell_text = []
                        for cell_elem in cell_content:
                            if 'paragraph' in cell_elem:
                                paragraph = cell_elem['paragraph']
                                elements = paragraph.get('elements', [])
                                for elem in elements:
                                    if 'textRun' in elem:
                                        text_content = elem['textRun'].get('content', '')
                                        cell_text.append(text_content)
                        text_parts.append(''.join(cell_text))
        
        return '\n'.join(text_parts)
    
    def upload_document_to_drive(self, file_content: str, filename: str, parent_folder_id: str = None) -> Optional[str]:
        """Upload a document to Google Drive with detailed error logging"""
        max_retries = 3
        retry_delay = 2
        
        for attempt in range(max_retries):
            try:
                from googleapiclient.http import MediaIoBaseUpload
                
                logger.info(f"Upload attempt {attempt + 1} for: {filename}")
                
                # Create file metadata
                file_metadata = {
                    'name': filename,
                    'mimeType': 'application/vnd.google-apps.document'
                }
                
                if parent_folder_id:
                    file_metadata['parents'] = [parent_folder_id]
                    logger.debug(f"Upload target folder: {parent_folder_id}")
                
                # Convert text content to file-like object
                content_bytes = file_content.encode('utf-8')
                logger.debug(f"File content size: {len(content_bytes)} bytes")
                
                media = MediaIoBaseUpload(
                    io.BytesIO(content_bytes),
                    mimetype='text/plain',
                    resumable=True,
                    chunksize=5 * 1024 * 1024  # 5MB chunks
                )
                
                # Upload file with detailed logging
                logger.info(f"Starting upload to Google Drive: {filename}")
                file = self.service.files().create(
                    body=file_metadata,
                    media_body=media,
                    fields='id,name,webViewLink,parents'
                ).execute()
                
                logger.info(f"Upload successful - ID: {file.get('id')}, Name: {file.get('name')}")
                logger.info(f"View link: {file.get('webViewLink', 'Not available')}")
                return file
                
            except Exception as e:
                logger.error(f"Upload attempt {attempt + 1} failed for {filename}: {str(e)}")
                logger.error(f"Error type: {type(e).__name__}")
                logger.error(f"Error details: {getattr(e, 'content', 'No additional details')}")
                
                if attempt < max_retries - 1:
                    backoff_time = retry_delay * (2 ** attempt)
                    logger.info(f"Retrying upload in {backoff_time}s...")
                    time.sleep(backoff_time)
                else:
                    logger.error(f"Failed to upload {filename} after {max_retries} attempts")
                    return None
        
        return None
    
    def create_google_doc_with_links(self, content: str, title: str) -> Optional[str]:
        """Create a professionally formatted Google Doc with active hyperlinks"""
        import re
        
        try:
            # Create a new document
            document = {
                'title': title
            }
            
            doc = self.docs_service.documents().create(body=document).execute()
            document_id = doc.get('documentId')
            
            # Format content for clean Google Docs display (remove markdown)
            formatted_content = GoogleDocsFormatter.format_content_for_google_docs(content)
            
            # Parse content and create requests for formatted content with hyperlinks
            requests = []
            
            # Split content into lines and process each line
            lines = formatted_content.split('\n')
            current_index = 1  # Start after the initial empty paragraph
            
            for line in lines:
                if line.strip():  # Only process non-empty lines
                    # Find URLs in the line
                    url_pattern = r'https?://[^\s<>"{}|\\^`[\]]+'
                    urls = re.findall(url_pattern, line)
                    
                    if urls:
                        # Process line with URLs - convert to hyperlinks
                        parts = re.split(url_pattern, line)
                        
                        for i, part in enumerate(parts):
                            if part:  # Add text part
                                requests.append({
                                    'insertText': {
                                        'location': {'index': current_index},
                                        'text': part
                                    }
                                })
                                current_index += len(part)
                            
                            # Add hyperlink if there's a corresponding URL
                            if i < len(urls):
                                url = urls[i]
                                requests.append({
                                    'insertText': {
                                        'location': {'index': current_index},
                                        'text': url
                                    }
                                })
                                
                                # Make it a hyperlink (let Google Docs use default link styling)
                                requests.append({
                                    'updateTextStyle': {
                                        'range': {
                                            'startIndex': current_index,
                                            'endIndex': current_index + len(url)
                                        },
                                        'textStyle': {
                                            'link': {'url': url}
                                        },
                                        'fields': 'link'
                                    }
                                })
                                current_index += len(url)
                    else:
                        # Simple text line without URLs
                        requests.append({
                            'insertText': {
                                'location': {'index': current_index},
                                'text': line
                            }
                        })
                        current_index += len(line)
                
                # Add line break
                requests.append({
                    'insertText': {
                        'location': {'index': current_index},
                        'text': '\n'
                    }
                })
                current_index += 1
            
            # Apply all formatting requests
            if requests:
                self.docs_service.documents().batchUpdate(
                    documentId=document_id,
                    body={'requests': requests}
                ).execute()
            
            # Get the document info to return web view link
            doc_info = self.service.files().get(
                fileId=document_id,
                fields='webViewLink,name'
            ).execute()
            
            logger.info(f"Successfully created Google Doc with hyperlinks: {title}")
            return {
                'id': document_id,
                'name': doc_info.get('name'),
                'webViewLink': doc_info.get('webViewLink')
            }
            
        except Exception as e:
            logger.error(f"Failed to create Google Doc with hyperlinks: {str(e)}")
            # Fallback to original upload method
            return self.upload_document_to_drive(content, title)

# =============================================================================
# DOCUMENT READER
# =============================================================================

class DocumentReader:
    """Handles reading different document formats"""
    
    @staticmethod
    def read_pdf(file_content: bytes) -> Optional[str]:
        """Extract text from PDF files - resilient to corrupt files"""
        try:
            if len(file_content) < 100:  # Too small to be a valid PDF
                return None
                
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            if len(pdf_reader.pages) == 0:
                return None
            
            text_content = []
            for i, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():  # Only add non-empty pages
                        text_content.append(page_text)
                except Exception as page_error:
                    # Skip corrupt pages but continue with others
                    continue
                    
                # Limit to first 50 pages to avoid memory issues
                if i >= 50:
                    break
            
            result = '\n'.join(text_content)
            return result if result.strip() else None
            
        except Exception as e:
            # Return None to try alternative methods
            return None
    
    @staticmethod
    def read_docx(file_content: bytes) -> Optional[str]:
        """Extract text from DOCX files - resilient to corrupt files"""
        try:
            if len(file_content) < 100:  # Too small to be valid DOCX
                return None
                
            docx_file = io.BytesIO(file_content)
            doc = DocxDocument(docx_file)
            
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Only add non-empty paragraphs
                    text_content.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)
            
            result = '\n'.join(text_content)
            return result if result.strip() else None
            
        except Exception as e:
            # Return None to try alternative methods
            return None
    
    @staticmethod
    def read_doc(file_content: bytes) -> Optional[str]:
        """Extract text from older DOC files"""
        try:
            # Try to extract text from older DOC format
            # This is more complex and may require additional libraries
            # For now, try to decode as plain text (limited effectiveness)
            try:
                # Attempt simple text extraction (not perfect but better than nothing)
                text = file_content.decode('utf-8', errors='ignore')
                # Clean up the text by removing non-printable characters
                import re
                cleaned_text = re.sub(r'[^\x20-\x7E\n\r\t]', '', text)
                return cleaned_text if cleaned_text.strip() else None
            except:
                return None
        except Exception as e:
            print(f"Error reading DOC: {str(e)}")
            return None
    
    @staticmethod
    def read_text(file_content: bytes) -> Optional[str]:
        try:
            # decode defensively
            return file_content.decode('utf-8', errors='replace')
        except Exception as e:
            print(f"Error reading text file: {str(e)}")
            return None
    
    @staticmethod
    def read_google_doc(file_content: str) -> Optional[str]:
        """Extract text from Google Docs (already extracted as text)"""
        try:
            # Google Docs content is already extracted as text by GoogleDriveClient
            return file_content
        except Exception as e:
            print(f"Error reading Google Doc: {str(e)}")
            return None
    
    @staticmethod
    def extract_urls(text: str) -> List[str]:
        """Extract URLs from text content"""
        if not text:
            return []
        
        # Enhanced URL regex pattern to catch various formats
        url_patterns = [
            r'https?://[^\s<>"{}|\\^`[\]]+',  # Standard HTTP/HTTPS URLs
            r'www\.[^\s<>"{}|\\^`[\]]+',      # www. URLs without protocol
            r'[a-zA-Z0-9][a-zA-Z0-9-]*\.(?:com|org|net|edu|gov|io|co|ai|tech|dev)[^\s<>"{}|\\^`[\]]*',  # Domain-based URLs
            r'(?:^|\s)(google\.com[^\s<>"{}|\\^`[\]]*)',  # Specific pattern for google.com
            r'[a-zA-Z0-9.-]+\.(?:com|org|net|edu|gov|io|co|ai|tech|dev)(?:/[^\s]*)?'  # More flexible domain matching
        ]
        
        urls = set()
        for pattern in url_patterns:
            found_urls = re.findall(pattern, text, re.IGNORECASE)
            for url in found_urls:
                # Clean and validate URL
                url = url.strip('.,;:!?"\'()[]{}')
                if len(url) > 10:  # Minimum reasonable URL length
                    # Add protocol if missing
                    if not url.startswith(('http://', 'https://')):
                        if url.startswith('www.'):
                            url = 'https://' + url
                        else:
                            url = 'https://' + url
                    urls.add(url)
        
        return list(urls)
    
    @staticmethod
    def validate_url(url: str) -> bool:
        """Validate if URL is properly formatted"""
        try:
            result = urllib.parse.urlparse(url)
            return all([result.scheme, result.netloc])
        except:
            return False
    
    @classmethod
    def extract_text(cls, file_content: bytes, mime_type: str) -> Optional[str]:
        """Extract text based on file type"""
        if mime_type == 'application/pdf':
            return cls.read_pdf(file_content)
        elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            # Modern DOCX format
            return cls.read_docx(file_content)
        elif mime_type == 'application/msword':
            # Older DOC format
            return cls.read_doc(file_content)
        elif mime_type == 'text/plain':
            return cls.read_text(file_content)
        elif mime_type == 'application/vnd.google-apps.document':
            # For Google Docs, the content should already be extracted as text
            # by GoogleDriveClient.get_file_content()
            if isinstance(file_content, str):
                return cls.read_google_doc(file_content)
            else:
                # If we get bytes, try to decode as text
                try:
                    text_content = file_content.decode('utf-8', errors='replace')
                    return cls.read_google_doc(text_content)
                except:
                    return None
        else:
            print(f"Unsupported file type: {mime_type}")
            return None
    
    @classmethod
    def get_document_info(cls, file_content, mime_type: str, file_name: str) -> Dict[str, Any]:
        """Get comprehensive document information including URLs"""
        # Handle both bytes (binary files) and str (Google Docs text)
        if isinstance(file_content, str):
            # This is already extracted text (Google Docs)
            text_content = file_content
        else:
            # This is binary content, extract text
            text_content = cls.extract_text(file_content, mime_type)
        
        if text_content is None:
            return {
                'file_name': file_name,
                'mime_type': mime_type,
                'content': None,
                'word_count': 0,
                'char_count': 0,
                'urls': [],
                'error': 'Failed to extract text'
            }
        
        # Extract URLs from the text content
        urls = cls.extract_urls(text_content)
        
        return {
            'file_name': file_name,
            'mime_type': mime_type,
            'content': text_content,
            'word_count': len(text_content.split()),
            'char_count': len(text_content),
            'urls': urls,
            'error': None
        }

# =============================================================================
# OPENAI LLM CLIENT
# =============================================================================

class OpenAIAnalyzer:
    """Handles OpenAI GPT interactions with web search via Responses API"""
    
    def __init__(self, api_key: str, model: str = "gpt-4o-mini", enable_web_search: bool = True,
                 system_prompt: Optional[str] = None, system_prompt_path: Optional[str] = None,
                 curriculum_manager: Optional[CurriculumManager] = None):
        self.client = OpenAI(api_key=api_key)
        self.model = model  # Note: Responses API works with gpt-4o-mini and newer models
        self.enable_web_search = enable_web_search
        self.curriculum_manager = curriculum_manager

        # Load system prompt
        self.system_prompt = system_prompt
        if not self.system_prompt and system_prompt_path and os.path.exists(system_prompt_path):
            try:
                with open(system_prompt_path, "r", encoding="utf-8") as f:
                    base_prompt = f.read()
                    
                # Use base system prompt without curriculum context for simple conversations
                # Curriculum context will be added conditionally when analyzing curriculum-related topics
                self.system_prompt = base_prompt
                logger.info(f"Loaded system prompt from {system_prompt_path} ({len(self.system_prompt)} characters)")
                    
            except Exception as e:
                logger.warning(f"Failed to load system prompt from {system_prompt_path}: {e}")
                self.system_prompt = None
        
        try:
            self.tokenizer = tiktoken.encoding_for_model(self.model)
        except KeyError:
            # Fallback to gpt-4 tokenizer for approximation if model not found
            self.tokenizer = tiktoken.encoding_for_model("gpt-4")
    
    def count_tokens(self, text: str) -> int:
        """Count tokens in text"""
        return len(self.tokenizer.encode(text))
    
    def generate_response_with_web_search(self, prompt: str) -> str:
        """Generate response using OpenAI Responses API with web search"""
        try:
            # sanitize prompt to ensure valid unicode
            prompt = sanitize_string(prompt)
            
            # Prepare input with system prompt if available
            input_payload = (
                [{"role": "system", "content": self.system_prompt},
                 {"role": "user", "content": prompt}]
                if self.system_prompt else prompt
            )
            
            if self.enable_web_search:
                try:
                    # Try to use the Responses API with web search
                    response = self.client.responses.create(
                        model=self.model,
                        tools=[{"type": "web_search_preview"}],
                        input=input_payload
                    )
                    return response.output_text
                except Exception as web_search_error:
                    logger.warning(f"Responses API with web search failed: {str(web_search_error)}")
                    # Check if it's a permission/authentication error
                    if "403" in str(web_search_error) or "forbidden" in str(web_search_error).lower():
                        logger.info("Web search requires special API permissions. Falling back to regular chat completion.")
                    else:
                        logger.info("Web search temporarily unavailable. Falling back to regular chat completion.")
                    
                    # Fallback to regular chat completions
                    messages = (
                        [{"role": "system", "content": self.system_prompt},
                         {"role": "user", "content": prompt}]
                        if self.system_prompt else [{"role": "user", "content": prompt}]
                    )
                    response = self.client.chat.completions.create(
                        model=self.model,
                        messages=messages,
                        max_tokens=4000,
                        temperature=0.1
                    )
                    return response.choices[0].message.content
            else:
                # Fallback to regular chat completions
                messages = (
                    [{"role": "system", "content": self.system_prompt},
                     {"role": "user", "content": prompt}]
                    if self.system_prompt else [{"role": "user", "content": prompt}]
                )
                response = self.client.chat.completions.create(
                    model=self.model,
                    messages=messages,
                    max_tokens=4000,
                    temperature=0.1
                )
                return response.choices[0].message.content
                
        except Exception as e:
            logger.error(f"OpenAI API error: {str(e)}")
            # Fallback to regular chat completions
            try:
                messages = (
                    [{"role": "system", "content": self.system_prompt},
                     {"role": "user", "content": prompt}]
                    if self.system_prompt else [{"role": "user", "content": prompt}]
                )
                response = self.client.chat.completions.create(
                    model=self.model,
                    messages=messages,
                    max_tokens=4000,
                    temperature=0.1
                )
                return response.choices[0].message.content
            except Exception as fallback_error:
                return f"Error: {str(e)}"
    
    def generate_response(self, prompt: str, max_tokens: int = 4000) -> str:
        """Generate response with optional web search"""
        return self.generate_response_with_web_search(prompt)
    
    def analyze_documents_with_web_research(self, documents: List[Dict[str, Any]], topic: str = None) -> str:
        """Analyze documents, URLs, and supplement with web research - returns comprehensive analysis"""
        
        # Extract URLs from all documents
        all_urls = []
        processed_docs = []
        url_analysis = ""
        
        if documents:
            processed_docs = [doc for doc in documents if doc.get('content')]
            
            # Collect all URLs from documents
            for doc in processed_docs:
                if doc.get('urls'):
                    all_urls.extend(doc['urls'])
            
            # Remove duplicates
            all_urls = list(set(all_urls))
            
            # Analyze URLs if found
            if all_urls:
                url_analysis = self._analyze_urls(all_urls[:10])  # Limit to first 10 URLs
        
        # Create comprehensive analysis prompt
        doc_summary = f"\nAnalyzing {len(processed_docs)} documents about {topic or 'the topic'}."
        url_summary = f"\nFound {len(all_urls)} URLs in documents." if all_urls else ""
        
        prompt = f"""
Run the 100x Engineers research workflow from the system prompt on this topic:

Topic: {topic or 'general analysis'}

{doc_summary}{url_summary}

Use the Drive excerpts and any extracted URLs as context. Search for information from BOTH recent sources (2024-2025) for latest updates AND historical sources (2020-2023) for foundational knowledge and trends. This balanced approach ensures current relevance while building proper context. Output strictly using the section headings defined in the system prompt's Output Format (TL;DR, First-Principles History, Cohort Relevance (Gate A), Teach Lane, Biz-Scout Lane, Community Sentiment, Risks & Unknowns, Bridge, Evidence & Citations, Reasoning Trace, Next Review Date & Watchlist). Use Asia/Kolkata timestamps. Cite official sources first, then community sentiment.

{url_analysis}

Apply the full evaluation framework including Gate A scoring, Teach/Biz-Scout lane analysis, and all quality checks.
        """
        
        # Use web search to get comprehensive information
        return self.generate_response_with_web_search(prompt)
    
    def _analyze_urls(self, urls: List[str]) -> str:
        """Analyze URLs found in documents using web search"""
        if not urls:
            return ""
        
        # Create prompt to analyze the URLs
        urls_list = "\n".join([f"- {url}" for url in urls[:10]])
        
        url_prompt = f"""
        Research and analyze these URLs found in the documents:
        {urls_list}
        
        For each relevant URL, provide:
        1. What type of resource it is
        2. Key information or insights from the content
        3. How it relates to the main topic
        
        Keep the analysis focused and relevant.
        """
        
        try:
            url_analysis = self.generate_response_with_web_search(url_prompt)
            return f"\n**URL Analysis from Documents:**\n{url_analysis}\n"
        except Exception as e:
            return f"\n**URLs found in documents:** {len(urls)} URLs identified.\n"
    
    def analyze_documents(self, documents: List[Dict[str, Any]], topic: str = None) -> str:
        """Analyze multiple documents and generate comprehensive report (legacy method)"""
        if not documents:
            return "No documents provided for analysis."
        
        document_summaries = []
        total_tokens = 0
        
        for doc in documents:
            if doc['content'] is None:
                continue
            
            # Limit content to manage token usage
            content = doc['content'][:10000]
            tokens = self.count_tokens(content)
            total_tokens += tokens
            
            if total_tokens > 100000:  # Token limit safety
                break
            
            summary_prompt = f"""
            Analyze the following document and provide a concise summary:
            
            Document: {doc['file_name']}
            Content: {content}
            
            Please provide:
            1. Main topics covered
            2. Key findings or information
            3. Important details
            
            Summary:
            """
            summary_prompt = sanitize_string(summary_prompt)
            summary = self.generate_response(summary_prompt, max_tokens=500)
            document_summaries.append({
                'file_name': doc['file_name'],
                'summary': summary,
                'word_count': doc['word_count']
            })
        
        return self._generate_final_report(document_summaries, topic)
    
    def _create_comprehensive_analysis_prompt(self, documents: List[Dict[str, Any]], topic: str = None) -> str:
        """Create a comprehensive prompt for document analysis with web research"""
        
        # Document summary section
        doc_summary = ""
        if documents:
            doc_summaries = []
            total_tokens = 0
            
            for doc in documents:
                if doc['content'] is None:
                    continue
                
                # Limit content to manage token usage
                content = doc['content'][:8000]  # Reduced for Responses API
                tokens = self.count_tokens(content)
                total_tokens += tokens
                
                if total_tokens > 80000:  # Conservative limit for Responses API
                    break
                
                doc_summaries.append(f"""
Document: {doc['file_name']}
Word Count: {doc.get('word_count', 0)}
Content Preview: {content}
                """)
            
            doc_summary = "\n---\n".join(doc_summaries)
        
        # Create comprehensive prompt
        prompt = f"""
Please provide a comprehensive analysis combining the document information below with current web research.

Topic Focus: {topic or 'General Document Analysis'}

DOCUMENT ANALYSIS:
{doc_summary if doc_summary else "No documents provided - focus on web research only."}

INSTRUCTIONS:
1. Analyze the provided documents (if any) for key themes and insights
2. Use web search to find current, relevant information about the topic
3. Compare document findings with current web information
4. Provide a structured report with:
   - Executive Summary
   - Key Findings from Documents (if any)
   - Current Web Research Insights
   - Comparative Analysis
   - Conclusions and Recommendations
   - Sources and References

Focus on providing factual, well-researched information from credible sources.
"""
        
        return prompt
    
    def _generate_final_report(self, summaries: List[Dict], topic: str = None) -> str:
        """Generate final comprehensive report"""
        summaries_text = "\n\n".join([
            f"Document: {s['file_name']}\nWord Count: {s['word_count']}\nSummary: {s['summary']}"
            for s in summaries
        ])
        
        topic_instruction = f"Focus specifically on information related to: {topic}" if topic else "Provide a comprehensive overview"
        
        final_prompt = f"""
        Based on the following document summaries from Google Drive, generate a comprehensive report.
        
        {topic_instruction}
        
        Document Summaries:
        {summaries_text}
        
        Please provide:
        1. Executive Summary
        2. Key Findings
        3. Main Topics Covered
        4. Important Insights
        5. Recommendations (if applicable)
        
        Report:
        """
        
        return self.generate_response(final_prompt, max_tokens=2000)
    
    def web_research_only(self, topic: str) -> str:
        """Perform web research only - returns brief useful data"""
        
        prompt = f"""
        Research topic: {topic}
        
        Search for information from BOTH recent sources (2024-2025) for current updates AND historical sources (2000-2023) for foundational trends and context. This balanced approach ensures comprehensive understanding.
        
        Provide brief, useful insights about {topic} including:
        1. Latest news and developments (2024-2025)
        2. Key trends and updates from recent years
        3. Important facts and historical context
        4. Foundational developments that shaped the field
        
        Keep it concise and actionable.
        """
        
        return self.generate_response_with_web_search(prompt)
    
    def answer_question(self, documents: List[Dict[str, Any]], question: str) -> str:
        """Answer specific questions about documents"""
        if not documents:
            return "No documents available to answer the question."
        
        relevant_content = []
        total_tokens = 0
        
        for doc in documents:
            if doc['content'] is None:
                continue
            
            content = doc['content'][:5000]  # Limit content
            tokens = self.count_tokens(content)
            total_tokens += tokens
            
            if total_tokens > 50000:  # Token limit
                break
            
            relevant_content.append(f"Document: {doc['file_name']}\nContent: {content}\n")
        
        context = "\n---\n".join(relevant_content)
        
        answer_prompt = f"""
        Based on the following documents from Google Drive, please answer this question:
        
        Question: {question}
        
        Documents:
        {context}
        
        Please provide a detailed answer based on the available information:
        """
        
        return self.generate_response(answer_prompt, max_tokens=1500)

# =============================================================================
# REPORT GENERATOR
# =============================================================================

class ReportGenerator:
    """Handles report generation and formatting"""
    
    def __init__(self, gdrive_client: Optional['GoogleDriveClient'] = None):
        self.console = Console()
        self.gdrive_client = gdrive_client
        self.downloads_dir = r"C:\Users\visha\Downloads"
        self.docx_generator = DocxReportGenerator(self.downloads_dir)
    
    def generate_report(
        self, 
        documents: List[Dict[str, Any]], 
        llm_analysis: str, 
        topic: str = None,
        save_to_file: bool = True,
        upload_to_drive: bool = True,
        model_name: str = "GPT-4o-mini"
    ) -> str:
        """Generate comprehensive report"""
        # Use Asia/Kolkata timezone
        kolkata_tz = pytz.timezone('Asia/Kolkata')
        timestamp = datetime.now(kolkata_tz).strftime("%Y-%m-%d %H:%M:%S %Z")
        
        # Create clean, professional report header
        header = f"""# {topic or 'Google Drive Analysis'} - 100x Engineers Report

**Generated:** {timestamp}  
**Analysis Model:** {model_name}  
**Research Methodology:** Google Drive + Web Search + 100x Engineers Framework"""
        
        # Create document summary
        doc_summary = self._create_document_summary(documents)
        
        
        # Combine all sections
        full_report = f"{header}\n\n{doc_summary}\n\n## Analysis Results\n\n{llm_analysis}"
        
        if save_to_file:
            filepath = self._save_report_to_file(full_report, topic)
            
            if upload_to_drive and self.gdrive_client:
                try:
                    uploaded_file = self.gdrive_client.create_google_doc_with_links(
                        full_report, 
                        os.path.basename(filepath).replace('.md', '')
                    )
                    if uploaded_file:
                        self.console.print(f"SUCCESS: Report uploaded to Google Drive: {uploaded_file.get('webViewLink', 'Link not available')}")
                    else:
                        self.console.print("⚠️ Failed to upload report to Google Drive")
                except Exception as e:
                    self.console.print(f"⚠️ Upload to Google Drive failed: {str(e)}")
        
        return full_report
    
    def _create_document_summary(self, documents: List[Dict[str, Any]]) -> str:
        """Create simplified document summary section"""
        if not documents:
            return "## Sources Analyzed\n\nNo documents were processed.\n\n"
        
        # Calculate statistics
        total_words = sum(doc.get('word_count', 0) for doc in documents if doc.get('content'))
        processed_docs = len([doc for doc in documents if doc.get('content')])
        failed_docs = len([doc for doc in documents if doc.get('error')])
        
        # Create clean, presentable summary
        summary = "## Sources Analyzed\n\n"
        summary += f"**Research Base:** {processed_docs} Google Drive documents ({total_words:,} words analyzed)"
        
        if failed_docs > 0:
            summary += f" • {failed_docs} documents could not be processed"
        
        summary += "\n\n"
        
        return summary
    
    def _save_report_to_file(self, report: str, topic: str = None) -> str:
        """Save report as .docx file to Downloads directory"""
        try:
            # Use DocxReportGenerator to create professional Word document
            topic_clean = topic.replace(' ', '_') if topic else "Report"
            title = f"Google Drive Analysis - {topic}" if topic else "Google Drive Analysis Report"
            
            filepath = self.docx_generator.create_docx_report(report, title)
            self.console.print(f"Report saved to: {filepath}")
            return filepath
            
        except Exception as e:
            # Fallback to plain text file if docx creation fails
            kolkata_tz = pytz.timezone('Asia/Kolkata')
            timestamp = datetime.now(kolkata_tz).strftime("%Y%m%d_%H%M%S")
            topic_str = f"_{topic.replace(' ', '_')}" if topic else ""
            filename = f"gdrive_analysis_report{topic_str}_{timestamp}.txt"
            
            fallback_dir = "reports"
            if not os.path.exists(fallback_dir):
                os.makedirs(fallback_dir)
            filepath = os.path.join(fallback_dir, filename)
            
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(report)
            self.console.print(f"Warning: Report saved to: {filepath} (fallback format)")
            return filepath
    
    def generate_enhanced_report(
        self, 
        documents: List[Dict[str, Any]], 
        llm_analysis: str, 
        topic: str = None,
        web_sources: List[str] = None,
        save_to_file: bool = True,
        upload_to_drive: bool = True,
        model_name: str = "GPT-4o-mini with Web Search"
    ) -> str:
        """Generate enhanced report with web research sources"""
        # Use Asia/Kolkata timezone
        kolkata_tz = pytz.timezone('Asia/Kolkata')
        timestamp = datetime.now(kolkata_tz).strftime("%Y-%m-%d %H:%M:%S %Z")
        
        # Create clean, professional report header
        header = f"""# {topic or 'Enhanced Google Drive Analysis'} - 100x Engineers Report

**Generated:** {timestamp}  
**Analysis Model:** {model_name}  
**Research Methodology:** Google Drive + Web Search + 100x Engineers Framework"""
        
        # Create document summary
        doc_summary = self._create_document_summary(documents)
        
        # Create web sources section
        web_sources_section = ""
        if web_sources:
            web_sources_section = f"""
## Web Research Sources

The following sources were consulted during the analysis:

{chr(10).join([f"- {source}" for source in web_sources[:20]])}
{"" if len(web_sources) <= 20 else f"... and {len(web_sources) - 20} more sources"}
            """
        
        # Combine all sections
        full_report = f"{header}\n\n{doc_summary}\n\n## Enhanced Analysis Results\n\n{llm_analysis}\n{web_sources_section}"
        
        if save_to_file:
            filepath = self._save_report_to_file(full_report, f"{topic}_enhanced" if topic else "enhanced")
            
            if upload_to_drive and self.gdrive_client:
                try:
                    uploaded_file = self.gdrive_client.create_google_doc_with_links(
                        full_report, 
                        f"Enhanced_Analysis_{topic or 'Report'}_{datetime.now().strftime('%Y%m%d')}"
                    )
                    if uploaded_file:
                        self.console.print("Enhanced report uploaded to Google Drive")
                    else:
                        self.console.print("Failed to upload enhanced report to Google Drive")
                except Exception as e:
                    self.console.print(f"Upload to Google Drive failed: {str(e)}")
        
        return full_report
    
    def display_report_summary(self, documents: List[Dict[str, Any]], topic: str = None):
        """Display report summary in console"""
        table = Table(title="Document Analysis Summary")
        table.add_column("Metric", style="cyan")
        table.add_column("Value", style="green")
        
        processed_docs = len([doc for doc in documents if doc.get('content')])
        failed_docs = len([doc for doc in documents if doc.get('error')])
        total_words = sum(doc.get('word_count', 0) for doc in documents if doc.get('content'))
        
        table.add_row("Topic", topic or "General Analysis")
        table.add_row("Total Documents", str(len(documents)))
        table.add_row("Successfully Processed", str(processed_docs))
        table.add_row("Failed to Process", str(failed_docs))
        table.add_row("Total Words", f"{total_words:,}")
        
        self.console.print(table)
        
        if failed_docs > 0:
            failed_panel = Panel(
                f"⚠️  {failed_docs} documents failed to process. Check the full report for details.",
                title="Warning",
                border_style="yellow"
            )
            self.console.print(failed_panel)
    
    def show_progress(self, total_files: int):
        """Create progress tracker"""
        return Progress(
            SpinnerColumn(),
            TextColumn("[progress.description]{task.description}"),
            console=self.console,
            transient=True
        )
    
    def create_question_response_report(
        self, 
        question: str, 
        answer: str, 
        documents: List[Dict[str, Any]],
        save_to_file: bool = True,
        upload_to_drive: bool = True,
        model_name: str = "GPT-4o-mini"
    ) -> str:
        """Create Q&A report"""
        # Use Asia/Kolkata timezone
        kolkata_tz = pytz.timezone('Asia/Kolkata')
        timestamp = datetime.now(kolkata_tz).strftime("%Y-%m-%d %H:%M:%S %Z")
        
        report = f"""
# Google Drive Q&A Report
Generated on: {timestamp}
Total Documents Consulted: {len(documents)}
LLM Model: {model_name}

## Question
{question}

## Answer
{answer}

## Source Documents
"""
        
        for i, doc in enumerate(documents, 1):
            if doc.get('content'):
                report += f"{i}. {doc['file_name']} ({doc.get('word_count', 0)} words)\n"
        
        if save_to_file:
            timestamp_file = datetime.now(kolkata_tz).strftime("%Y%m%d_%H%M%S")
            filename = f"gdrive_qa_report_{timestamp_file}.md"
            
            # Save to Downloads directory
            if not os.path.exists(self.downloads_dir):
                os.makedirs(self.downloads_dir)
            
            filepath = os.path.join(self.downloads_dir, filename)
            
            try:
                with open(filepath, 'w', encoding='utf-8') as f:
                    f.write(report)
                self.console.print(f"Q&A Report saved to: {filepath}")
                
                if upload_to_drive and self.gdrive_client:
                    try:
                        uploaded_file = self.gdrive_client.create_google_doc_with_links(
                            report, 
                            filename.replace('.md', '')
                        )
                        if uploaded_file:
                            self.console.print(f"SUCCESS: Q&A Report uploaded to Google Drive: {uploaded_file.get('webViewLink', 'Link not available')}")
                        else:
                            self.console.print("⚠️ Failed to upload Q&A report to Google Drive")
                    except Exception as e:
                        self.console.print(f"⚠️ Upload to Google Drive failed: {str(e)}")
                        
            except Exception as e:
                # Fallback to current directory
                fallback_dir = "reports"
                if not os.path.exists(fallback_dir):
                    os.makedirs(fallback_dir)
                filepath = os.path.join(fallback_dir, filename)
                with open(filepath, 'w', encoding='utf-8') as f:
                    f.write(report)
                self.console.print(f"⚠️ Q&A Report saved to fallback location: {filepath}")
        
        return report

# =============================================================================
# COMMAND LINE INTERFACE
# =============================================================================

console = Console()

@click.group()
@click.version_option(version="1.0.0", prog_name="GDrive LLM Analyzer")
def cli():
    """Google Drive Document Analysis Tool with OpenAI GPT-4o-mini"""
    pass

@cli.command()
@click.option('--topic', '-t', help='Specific topic to focus on in the analysis')
@click.option('--query', '-q', help='Google Drive search query to filter documents')
@click.option('--limit', '-l', type=int, default=10,help='Maximum number of documents to analyze')
@click.option('--interactive', '-i', is_flag=True, help='Interactive mode for report review')
@click.option('--web-research/--no-web-research', default=True, help='Enable/disable web research (default: enabled)')
@click.option('--docs-only', is_flag=True, help='Analyze only documents without web research')
def analyze(topic: Optional[str], query: Optional[str], limit: int, interactive: bool, web_research: bool, docs_only: bool):
    """Analyze documents from Google Drive with optional web research using OpenAI's Responses API"""
    
    # If docs_only is specified, disable web research
    if docs_only:
        web_research = False
    
    try:
        config = Config()
        
        title = "Google Drive Analysis" + (" + Web Research" if web_research else " (Documents Only)")
        console.print(Panel(title, style="blue"))
        
        # Initialize Google Drive client
        console.print("Connecting to Google Drive...")
        gdrive_client = GoogleDriveClient(
            credentials_path=config.google_credentials_path,
            drive_token_path=config.google_drive_token_path,
            docs_token_path=config.google_docs_token_path
        )
        
        # Initialize report generator with gdrive_client
        report_gen = ReportGenerator(gdrive_client)
        
        # Get list of documents
        console.print(f"Searching for documents... (limit: {limit})")
        files = gdrive_client.list_files(query=query)
        
        if not files and not web_research:
            console.print("No documents found and web research is disabled.")
            return
        
        files = files[:limit] if files else []
        if files:
            console.print(f"Found {len(files)} documents")
        
        # Process documents with batch processing optimization
        documents = []
        if files:
            batch_size = min(2, len(files))  # Process up to 2 files concurrently for better stability
            
            def process_single_file(file_info):
                """Process a single file and return document info"""
                try:
                    # Handle Google Docs specially
                    if file_info['mimeType'] == 'application/vnd.google-apps.document':
                        # Use Google Docs API for better content extraction
                        text_content = gdrive_client.get_file_content(file_info['id'])
                        if text_content:
                            doc_info = DocumentReader.get_document_info(
                                text_content,  # Pass as string
                                file_info['mimeType'], 
                                file_info['name']
                            )
                            logger.info(f"Successfully processed Google Doc: {file_info['name']} ({doc_info.get('word_count', 0)} words)")
                        else:
                            doc_info = {
                                'file_name': file_info['name'],
                                'mime_type': file_info['mimeType'],
                                'content': None,
                                'word_count': 0,
                                'char_count': 0,
                                'urls': [],
                                'error': 'Failed to extract Google Docs content'
                            }
                            logger.warning(f"Failed to extract content from Google Doc: {file_info['name']}")
                    else:
                        # Handle binary files (PDF, DOCX, etc.)
                        file_content = gdrive_client.download_file(file_info['id'], file_info['name'])
                        
                        if file_content:
                            doc_info = DocumentReader.get_document_info(
                                file_content,  # Pass as bytes
                                file_info['mimeType'], 
                                file_info['name']
                            )
                            logger.info(f"Successfully processed binary file: {file_info['name']} ({doc_info.get('word_count', 0)} words)")
                        else:
                            doc_info = {
                                'file_name': file_info['name'],
                                'mime_type': file_info['mimeType'],
                                'content': None,
                                'word_count': 0,
                                'char_count': 0,
                                'urls': [],
                                'error': 'Failed to download file'
                            }
                            logger.warning(f"Failed to download binary file: {file_info['name']}")
                    
                    # Add small delay between file processing to avoid overwhelming the API
                    time.sleep(0.1)
                    
                    return doc_info
                    
                except Exception as e:
                    logger.error(f"Error processing file {file_info['name']}: {str(e)}")
                    return {
                        'file_name': file_info['name'],
                        'mime_type': file_info['mimeType'],
                        'content': None,
                        'word_count': 0,
                        'char_count': 0,
                        'urls': [],
                        'error': f'Processing error: {str(e)}'
                    }
            
            with report_gen.show_progress(len(files)) as progress:
                task = progress.add_task("Processing documents...", total=len(files))
                
                # Process files in batches using ThreadPoolExecutor
                with ThreadPoolExecutor(max_workers=batch_size) as executor:
                    # Submit all tasks
                    future_to_file = {executor.submit(process_single_file, file_info): file_info for file_info in files}
                    
                    # Collect results as they complete
                    for future in as_completed(future_to_file):
                        file_info = future_to_file[future]
                        try:
                            doc_info = future.result()
                            documents.append(doc_info)
                            progress.update(task, description=f"Processed {file_info['name']}")
                            progress.advance(task)
                        except Exception as e:
                            logger.error(f"Failed to process {file_info['name']}: {str(e)}")
                            progress.advance(task)
            
            # Display processing summary
            report_gen.display_report_summary(documents, topic)
        
        # Initialize OpenAI analyzer with Responses API
        console.print("Initializing OpenAI Responses API...")
        # Use gpt-4o-mini for Responses API compatibility
        model = "gpt-4o-mini" if web_research else config.openai_model
        analyzer = OpenAIAnalyzer(
            config.openai_api_key, 
            model, 
            enable_web_search=web_research,
            system_prompt_path=config.system_prompt_path,
            curriculum_manager=config.curriculum_manager
        )
        
        # Perform analysis
        if web_research:
            analysis_result = analyzer.analyze_documents_with_web_research(documents, topic)
        else:
            analysis_result = analyzer.analyze_documents(documents, topic)
        
        # Display result to console
        console.print("\n" + "="*80)
        console.print("ANALYSIS RESULTS:")
        console.print("="*80)
        console.print(analysis_result)
        
        # Generate and save report with upload
        console.print("\nGenerating report...")
        if web_research:
            full_report = report_gen.generate_enhanced_report(
                documents, 
                analysis_result, 
                topic,
                save_to_file=True,
                upload_to_drive=True,
                model_name=f"{analyzer.model} with Web Search"
            )
        else:
            full_report = report_gen.generate_report(
                documents, 
                analysis_result, 
                topic,
                save_to_file=True,
                upload_to_drive=True,
                model_name=analyzer.model
            )
            
    except Exception as e:
        console.print(f"ERROR: {str(e)}", style="red")

@cli.command()
@click.argument('question')
@click.option('--query', '-q', help='Google Drive search query to filter documents')
@click.option('--limit', '-l', type=int, default=10,help='Maximum number of documents to search')
def ask(question: str, query: Optional[str], limit: int):
    """Ask a question about your Google Drive documents"""
    
    try:
        config = Config()
        
        console.print(Panel(f"Question: {question}", style="blue"))
        
        # Initialize Google Drive client
        console.print("Connecting to Google Drive...")
        gdrive_client = GoogleDriveClient(
            credentials_path=config.google_credentials_path,
            drive_token_path=config.google_drive_token_path,
            docs_token_path=config.google_docs_token_path
        )
        
        # Initialize report generator with gdrive_client
        report_gen = ReportGenerator(gdrive_client)
        
        # Get list of documents
        console.print(f"Searching for relevant documents... (limit: {limit})")
        files = gdrive_client.list_files(query=query)
        
        if not files:
            console.print("No documents found.")
            return
        
        files = files[:limit]
        console.print(f"Found {len(files)} documents")
        
        # Process documents
        documents = []
        with report_gen.show_progress(len(files)) as progress:
            task = progress.add_task("Processing documents...", total=len(files))
            
            for file_info in files:
                progress.update(task, description=f"Processing {file_info['name']}")
                
                # Download and process file
                file_content = gdrive_client.download_file(file_info['id'], file_info['name'])
                
                if file_content:
                    doc_info = DocumentReader.get_document_info(
                        file_content, 
                        file_info['mimeType'], 
                        file_info['name']
                    )
                else:
                    # Try to get plain text content
                    text_content = gdrive_client.get_file_content(file_info['id'])
                    if text_content:
                        doc_info = {
                            'file_name': file_info['name'],
                            'mime_type': file_info['mimeType'],
                            'content': text_content,
                            'word_count': len(text_content.split()),
                            'char_count': len(text_content),
                            'error': None
                        }
                    else:
                        doc_info = {
                            'file_name': file_info['name'],
                            'mime_type': file_info['mimeType'],
                            'content': None,
                            'word_count': 0,
                            'char_count': 0,
                            'error': 'Failed to download or read file'
                        }
                
                documents.append(doc_info)
                progress.advance(task)
        
        # Initialize OpenAI analyzer
        console.print("AI Searching for answer with GPT-4o-mini...")
        analyzer = OpenAIAnalyzer(
            config.openai_api_key, 
            config.openai_model,
            system_prompt_path=config.system_prompt_path,
            curriculum_manager=config.curriculum_manager
        )
        
        # Get answer
        answer = analyzer.answer_question(documents, question)
        
        # Display answer
        console.print(Panel(answer, title="Answer", border_style="green"))
        
        # Generate Q&A report
        qa_report = report_gen.create_question_response_report(question, answer, documents, upload_to_drive=True, model_name=analyzer.model)
        console.print("Q&A report saved to Downloads directory and uploaded to Google Drive")
        
    except Exception as e:
        console.print(f"ERROR: {str(e)}", style="red")

@cli.command()
@click.argument('topic')
@click.option('--save-report', '-s', is_flag=True, help='Save research report to file')
@click.option('--format', '-f', type=click.Choice(['md', 'txt', 'json']), default='md', help='Output format')
def research(topic: str, save_report: bool, format: str):
    """Perform web research using OpenAI's Responses API"""
    
    try:
        config = Config()
        console.print(Panel(f"Web Research: {topic}", style="blue"))
        
        # Initialize analyzer with web search
        analyzer = OpenAIAnalyzer(
            config.openai_api_key, 
            "gpt-4o-mini", 
            enable_web_search=True,
            system_prompt_path=config.system_prompt_path,
            curriculum_manager=config.curriculum_manager
        )
        
        # Perform research and display brief results
        result = analyzer.web_research_only(topic)
        console.print(result)
        
        # Always save research result and upload to Drive
        kolkata_tz = pytz.timezone('Asia/Kolkata')
        timestamp = datetime.now(kolkata_tz).strftime('%Y-%m-%d %H:%M:%S %Z')
        full_research_report = f"# Web Research: {topic}\n\nGenerated on: {timestamp}\nModel: gpt-4o-mini with Web Search\n\n{result}"
        
        # Initialize report generator and save/upload
        try:
            gdrive_client = GoogleDriveClient(
                credentials_path=config.google_credentials_path,
                drive_token_path=config.google_drive_token_path,
                docs_token_path=config.google_docs_token_path
            )
            report_gen = ReportGenerator(gdrive_client)
            
            # Save locally to Downloads
            timestamp = datetime.now(kolkata_tz).strftime("%Y%m%d_%H%M%S")
            topic_clean = topic.replace(' ', '_').replace('/', '-')
            filename = f"web_research_{topic_clean}_{timestamp}.md"
            
            downloads_dir = r"C:\Users\visha\Downloads"
            if not os.path.exists(downloads_dir):
                os.makedirs(downloads_dir)
            
            filepath = os.path.join(downloads_dir, filename)
            
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(full_research_report)
            
            console.print(f"Report saved to: {filepath}")
            
            # Upload to Google Drive
            uploaded_file = gdrive_client.create_google_doc_with_links(
                full_research_report, 
                f"Web_Research_{topic_clean}_{timestamp}"
            )
            
            if uploaded_file:
                console.print(f"Report uploaded to Google Drive: {uploaded_file.get('webViewLink', 'Link not available')}")
            else:
                console.print("Failed to upload report to Google Drive")
                
        except Exception as upload_error:
            console.print(f"Google Drive upload failed: {str(upload_error)}", style="red")
            # Still save locally as fallback
            fallback_dir = "reports"
            os.makedirs(fallback_dir, exist_ok=True)
            filepath = os.path.join(fallback_dir, filename)
            
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(full_research_report)
            
            console.print(f"Saved to fallback location: {filepath}")
        
        if save_report:
            # For backwards compatibility, also save in old format if explicitly requested
            kolkata_tz = pytz.timezone('Asia/Kolkata')
            timestamp = datetime.now(kolkata_tz).strftime("%Y%m%d_%H%M%S")
            topic_clean = topic.replace(' ', '_').replace('/', '-')
            filename_old = f"web_research_{topic_clean}_{timestamp}.{format}"
            
            os.makedirs("reports", exist_ok=True)
            filepath_old = os.path.join("reports", filename_old)
            
            with open(filepath_old, 'w', encoding='utf-8') as f:
                f.write(full_research_report)
            
            console.print(f"Also saved in old format to: {filepath_old}")
            
    except Exception as e:
        console.print(f"ERROR: {str(e)}", style="red")


@cli.command("voice")
@click.option("--dev", is_flag=True, help="Run in dev mode")
def voice_cmd(dev: bool):
    """
    Start the LiveKit voice worker - simplified version
    """
    if not LIVEKIT_AVAILABLE:
        console.print("[red]LiveKit not available[/red]")
        return
    
    console.print("[cyan]Starting LiveKit worker...[/cyan]")
    console.print(f"[dim]URL: {os.getenv('LIVEKIT_URL')}[/dim]")
    console.print(f"[dim]KEY: {'***' if os.getenv('LIVEKIT_API_KEY') else 'Not set'}[/dim]")
    
    try:
        # SIMPLIFIED: No request_fnc, no custom port/host
        worker_opts = WorkerOptions(
            entrypoint_fnc=livekit_entrypoint
            # That's it! No other parameters
        )
        
        sys.argv = [sys.argv[0], "dev" if dev else "start"]
        lk_cli.run_app(worker_opts)
        
    except Exception as e:
        console.print(f"[red]Error: {e}[/red]")
        import traceback
        traceback.print_exc()

@cli.command("console")
def console_cmd():
    """
    Start the local text console (no LiveKit).
    Usage: python gdrive_analyzer.py console
    """
    ConversationalGDriveAnalyzer(voice_mode=False).run()


@cli.command()
def setup():
    """Setup configuration for the tool"""
    
    console.print(Panel("Setup: Google Drive LLM Analyzer Setup", style="blue"))
    
    # Create .env file if it doesn't exist
    env_content = """# Google Drive API credentials
GOOGLE_CREDENTIALS_PATH=credentials.json

# OpenAI API Configuration
OPENAI_API_KEY=your_openai_api_key_here
OPENAI_MODEL=gpt-4o-mini
"""
    
    if not os.path.exists(".env"):
        with open(".env", "w") as f:
            f.write(env_content)
        console.print("SUCCESS: Created .env file")
    
    console.print("\nConfiguration Steps:")
    console.print("1. Set up Google Drive API credentials:")
    console.print("   - Go to Google Cloud Console (https://console.cloud.google.com/)")
    console.print("   - Enable Google Drive API")
    console.print("   - Create OAuth 2.0 credentials (Desktop Application)")
    console.print("   - Download credentials.json file")
    console.print("   - Place it in the same directory as this script")
    
    console.print("\n2. Configure OpenAI API:")
    console.print("   - Get API key from OpenAI (https://platform.openai.com/api-keys)")
    console.print("   - Update .env file with your API key")
    
    console.print(f"\n3. Edit .env file and replace 'your_openai_api_key_here' with your actual API key")
    
    if Confirm.ask("Open .env file for editing?"):
        import subprocess
        try:
            subprocess.run(["notepad", ".env"], check=True)
        except:
            console.print("Please manually edit .env file")

@cli.command("test-voice-room")
@click.option("--room-name", default="test-gdrive", help="Room name")
def test_voice_room(room_name: str):
    """Create a test room to trigger the voice agent"""
    import asyncio
    from livekit import api
    
    async def create_and_join_room():
        try:
            # Get credentials
            url = os.getenv("LIVEKIT_URL", "ws://localhost:7880")
            api_key = os.getenv("LIVEKIT_API_KEY", "devkey")
            api_secret = os.getenv("LIVEKIT_API_SECRET", "secret")
            
            console.print(f"[cyan]Creating room '{room_name}' at {url}[/cyan]")
            
            # Create access token
            token = api.AccessToken(api_key, api_secret)
            token.with_identity("test-user")
            token.with_name("Test User")
            token.with_grants(api.VideoGrants(
                room_join=True,
                room=room_name,
                can_publish=True,
                can_subscribe=True
            ))
            
            jwt_token = token.to_jwt()
            
            console.print("[green]Room created successfully![/green]")
            console.print(f"[yellow]Room: {room_name}[/yellow]")
            console.print(f"[yellow]Token: {jwt_token[:50]}...[/yellow]")
            console.print(f"[yellow]URL: {url}[/yellow]")
            console.print("\n[cyan]Your agent should now be triggered if the worker is running![/cyan]")
            
            # Optional: Connect to the room programmatically
            from livekit import rtc
            room = rtc.Room()
            
            @room.on("participant_connected")
            def on_participant_connected(participant):
                console.print(f"[green]Participant connected: {participant.identity}[/green]")
            
            await room.connect(url, jwt_token)
            console.print("[green]Connected to room![/green]")
            
            # Keep connection alive
            await asyncio.sleep(60)  # Stay connected for 60 seconds
            
        except Exception as e:
            console.print(f"[red]Error creating room: {e}[/red]")
    
    asyncio.run(create_and_join_room())

@cli.command("voice-local")
def voice_local():
    """Run voice agent in local test mode (no LiveKit server needed)"""
    console.print("[cyan]Starting local voice test mode...[/cyan]")
    
    # This bypasses LiveKit and tests your agent directly
    async def test_agent():
        agent = VoiceGDriveAgent()
        
        # Simulate a voice command
        test_topic = "AI agents"
        console.print(f"[yellow]Testing analysis of: {test_topic}[/yellow]")
        
        # Create a mock context
        from types import SimpleNamespace
        mock_context = SimpleNamespace()
        
        # Test the analyze function directly
        result = await agent.analyze_drive_topic(
            mock_context, 
            topic=test_topic,
            query="AI agents curriculum",
            docs_only=True
        )
        
        console.print(f"[green]Result:[/green] {result}")
    
    import asyncio
    asyncio.run(test_agent())


@cli.command()
def test():
    """Test connections to Google Drive and OpenAI"""
    
    try:
        config = Config()
        console.print(Panel("Testing Connections", style="blue"))
        
        # Test Google Drive connection
        console.print("Testing Google Drive connection...")
        try:
            gdrive_client = GoogleDriveClient(
                credentials_path=config.google_credentials_path,
                drive_token_path=config.google_drive_token_path,
                docs_token_path=config.google_docs_token_path
            )
            files = gdrive_client.list_files()
            console.print(f"Google Drive: Connected ({len(files)} files accessible)")
        except Exception as e:
            console.print(f"ERROR Google Drive: {str(e)}", style="red")
        
        # Test OpenAI connection
        console.print("Testing OpenAI connection...")
        try:
            analyzer = OpenAIAnalyzer(
                config.openai_api_key, 
                config.openai_model,
                system_prompt_path=config.system_prompt_path,
                curriculum_manager=config.curriculum_manager
            )
            test_response = analyzer.generate_response("Hello, this is a test. Please respond with 'Test successful!'", max_tokens=50)
            console.print(f"OpenAI (gpt-4o-mini): {test_response}")
        except Exception as e:
            console.print(f"ERROR OpenAI: {str(e)}", style="red")
            
    except Exception as e:
        console.print(f"ERROR Configuration: {str(e)}", style="red")

# =============================================================================
# MAIN ENTRY POINT
# =============================================================================
# =============================================================================
# MAIN ENTRY POINT
# =============================================================================
# --- Unified LiveKit entrypoint (replace other entrypoint defs) ---

log = logging.getLogger("gdrive.livekit.entrypoint")

async def livekit_entrypoint(*args, **kwargs):
    """
    Unified entrypoint that supports:
      - livekit calling with (ctx) where ctx.connect exists
      - or livekit calling with (session) where session.start exists
    This logs immediately so we can see whether it was invoked and which branch runs.
    """
    try:
        log.info("[livekit_entrypoint] invoked with args=%s kwargs=%s", args and [type(a).__name__ for a in args] or [], list(kwargs.keys()))
        print("[livekit_entrypoint] invoked")  # immediate stdout marker

        # determine if first arg is a ctx-like (has connect) or session-like (has start)
        first = args[0] if args else kwargs.get("ctx") or kwargs.get("session")
        # branch A: ctx-like object (has connect)
        if first is not None and hasattr(first, "connect"):
            ctx = first
            log.info("[livekit_entrypoint] detected ctx-style invocation")
            print("[livekit_entrypoint] detected ctx-style invocation")
            # use same flow you had earlier that connected and built an AgentSession
            await ctx.connect(auto_subscribe="AUDIO_ONLY")

            session = AgentSession(
                stt=lk_openai.STT(model="gpt-4o-transcribe") if lk_openai else None,
                llm=lk_openai.LLM(model="gpt-4o-mini") if lk_openai else None,
                tts=lk_openai.TTS(
                    model="gpt-4o-mini-tts",
                    voice="ash",
                    instructions="Speak quickly and clearly; be concise and confident.",
                    speed=1.2,
                ) if lk_openai else None,
                vad=silero.VAD.load() if silero else None,
                turn_detection="vad",
            )

            agent = VoiceGDriveAgent()
            try:
                await agent.update_tools([agent.analyze_drive_topic])
                log.info("[livekit_entrypoint] tools registered (ctx-style)")
                print("[livekit_entrypoint] tools registered (ctx-style)")
            except Exception as e:
                log.warning("[livekit_entrypoint] update_tools warning (ctx-style): %s", e)
                print("[livekit_entrypoint] update_tools warning (ctx-style):", e)

            await session.start(
                room=ctx.room if hasattr(ctx, "room") else None,
                agent=agent,
                room_input_options=RoomInputOptions(
                    noise_cancellation=noise_cancellation.BVC() if noise_cancellation else None
                ),
            )

            await session.generate_reply(instructions="Hello — I can analyze your Drive and produce reports.")
            return

        # branch B: session-like object (has start)
        if first is not None and hasattr(first, "start"):
            session = first
            log.info("[livekit_entrypoint] detected session-style invocation")
            print("[livekit_entrypoint] detected session-style invocation")
            agent = VoiceGDriveAgent()
            try:
                print("[livekit_entrypoint] registering tools...")
                await agent.update_tools([agent.analyze_drive_topic])
                print("[livekit_entrypoint] tools registered.")
                log.info("[livekit_entrypoint] tools registered (session-style)")
            except Exception as e:
                print("[livekit_entrypoint] warning: update_tools failed:", repr(e))
                log.warning("[livekit_entrypoint] update_tools failed (session-style): %s", traceback.format_exc())

            # Some livekit versions expect `await session.start(agent=agent)`:
            try:
                await session.start(agent=agent)
                print("[livekit_entrypoint] session.start returned")
            except TypeError:
                # older/newer signature mismatch - try alternative
                try:
                    await session.start(room=None, agent=agent)
                    print("[livekit_entrypoint] session.start(room=None, agent=agent) returned")
                except Exception as e:
                    print("[livekit_entrypoint] session.start failed:", repr(e))
                    log.error("[livekit_entrypoint] session.start failed: %s", traceback.format_exc())
            return

        # If we get here, signature is unexpected
        log.error("[livekit_entrypoint] Could not detect ctx or session in args: %s %s", args, kwargs)
        print("[livekit_entrypoint] ERROR: Unknown invocation signature")
    except Exception:
        log.error("[livekit_entrypoint] Exception:\n" + traceback.format_exc())
        print("[livekit_entrypoint] Exception (see logs):")
        traceback.print_exc()


"""if __name__ == "__main__":
    import sys
    from dotenv import load_dotenv
    import os

    load_dotenv()

    # Basic diagnostics before attempting to start LiveKit
    def _voice_diagnostics():
        print("=== Voice diagnostics ===")
        print("LIVEKIT_AVAILABLE:", LIVEKIT_AVAILABLE)  # runtime import may still happen
        print("LIVEKIT_URL:", bool(os.getenv("LIVEKIT_URL")))
        print("LIVEKIT_API_KEY:", bool(os.getenv("LIVEKIT_API_KEY")))
        print("LIVEKIT_API_SECRET:", bool(os.getenv("LIVEKIT_API_SECRET")))
        print("=========================")

    if len(sys.argv) > 1 and sys.argv[1] == "voice":
        _voice_diagnostics()
        try:
            from livekit.agents import cli as lk_cli, WorkerOptions
        except Exception as e:
            print(f"LiveKit not available or import failed: {e}. Falling back to local console.")
            ConversationalGDriveAnalyzer(voice_mode=True).run()
            sys.exit(0)

        # Optional: forward extra args to LiveKit CLI by rewriting sys.argv if needed:
        # (uncomment if you want to support `python gdrive_analyzer.py voice start`)
        # extra = sys.argv[2:]
        # if extra:
        #     sys.argv = [sys.argv[0]] + extra
        # else:
        #     sys.argv = [sys.argv[0], "start"]

        # run the LiveKit worker (this blocks)
        sys.argv = [sys.argv[0], "start"]  # or "dev"
        lk_cli.run_app(WorkerOptions(entrypoint_fnc=livekit_entrypoint))

    elif len(sys.argv) > 1 and sys.argv[1] == "console":
        ConversationalGDriveAnalyzer().run()

    else:
        # Hand off to your Click commands (analyze / research / ask / setup / test)
        cli()"""


if __name__ == "__main__":
    load_dotenv()
    cli()




